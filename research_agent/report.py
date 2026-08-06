"""curation-report-synthesis Phase 4: generates the literature-review
report over a curation session's exact hand-picked paper set — a
different, 3-part structure (findings / limitations / future scope)
from summarize.py's theme-clustering, but reusing the identical
grounding TECHNIQUE: a dynamic Literal built from the exact paper_ids
in play, so the model is structurally unable to cite a paper outside
that set. summarize.py itself is untouched — this is a new, separate
schema/function, not a rewrite of it (see this module's own tests for
the adversarial proof that a report never cites a seen-but-rejected
candidate).

Deliberately a plain function, not a graph node — matching
summarize.py's own generate_summary()/generate_web_summary() pattern
exactly. By the time session.stage == "synthesize", the curation
DECISION is already finalized (that's what Phase 3's interactive loop
was for); there's nothing left here to pause/resume/branch on. One
topic + one fixed paper set -> one LLM call -> one structured result,
the same shape as summarize.py's own one-shot generation. A future
graph (e.g. Phase 5's chat) can call this plain function directly from
one of its own nodes, the same way qa.py's _generate_node already
calls the plain function _generate_answer internally.
"""

from __future__ import annotations

import io
import logging
import re
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Literal
from xml.sax.saxutils import escape as xml_escape, quoteattr as xml_quoteattr

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langfuse import get_client, observe
from openai import OpenAI
from pydantic import BaseModel, Field, create_model
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from research_agent.citations import format_apa_citation, format_web_citation
from research_agent.query_expansion import PaperPoolSession
from research_agent.schema import Paper, WebArticle
from research_agent.tracing import paper_metadata

logger = logging.getLogger(__name__)

# Same model choice/reasoning as summarize.py's SUMMARY_MODEL: infrequent
# (once per finished curation session, not looped), faithfulness matters
# more than cost for a call this rare.
REPORT_MODEL = "gpt-4.1"

# report-quality Phase R2C: which reader-depth template generated (or
# should generate) a report. Same plain-Literal-type convention as
# citations.py's own CitationStyle -- not a Python Enum class, matching
# how this codebase already represents "one of a small fixed set of
# strings" elsewhere (e.g. PaperPoolSession.stage's own "curate"/
# "synthesize"/"terminate" values). All three templates share the exact
# same 8 section keys/titles (see REPORT_TEMPLATES below) -- only prompt
# instructions and word-budget guidance vary per template.
ReportTemplate = Literal["foundational", "analytical", "expert"]
DEFAULT_REPORT_TEMPLATE: ReportTemplate = "analytical"

# report-quality Phase R4.1: whether/how the optional draft -> evaluate
# -> revise -> finalize refinement loop runs. Same plain-Literal-type
# convention as ReportTemplate above. "off" (the default) means every
# existing caller that doesn't opt in gets byte-identical behavior --
# refine_report_if_requested below returns its input completely
# untouched, zero extra LLM calls. "single" means at most ONE revision
# round -- see refine_report_if_requested's own docstring for the exact
# bounded flow. Deliberately a 2-value Literal, not a bool: R4.3's own
# planned configurable-depth values (e.g. a future "double") are
# additive to this vocabulary, not a breaking change to it, the same
# way ReportTemplate's own 3-value vocabulary was chosen over a bool.
RefinementMode = Literal["off", "single"]
DEFAULT_REFINEMENT_MODE: RefinementMode = "off"

FINDINGS_DESCRIPTION = "Synthesized findings and key contributions across the selected papers, grounded strictly in their abstracts."
LIMITATIONS_DESCRIPTION = "Limitations or shortcomings noted across the selected papers, and/or gaps or disagreements between them. State explicitly if none are apparent -- don't invent one."
FUTURE_SCOPE_DESCRIPTION = "Future research directions plausibly implied by the selected papers -- may be more speculative than the other two sections, but must still stay grounded in what the papers actually cover, not outside knowledge."

SECTION_NAMES = ("findings", "limitations", "future_scope")

SYSTEM_PROMPT = f"""You are writing a literature review report over a specific, deliberately hand-picked set of papers for a research topic. You will be given the topic and exactly this set of papers (id, title, abstract only) -- this is the FINAL set a user has already chosen; do not second-guess or add papers beyond it.

Write exactly three sections:
1. Findings: {FINDINGS_DESCRIPTION}
2. Limitations: {LIMITATIONS_DESCRIPTION}
3. Future scope: {FUTURE_SCOPE_DESCRIPTION}

For EACH section, ground every claim STRICTLY in the given abstracts -- do not use outside knowledge about these papers, their authors, or the topic beyond what the abstracts state. List which paper_ids actually support that section's content in its own cited_paper_ids, in the order you'd reference them. A section's citations are independent of the other two sections' -- a paper cited in Findings does not need to also appear in Limitations or Future scope, and vice versa.

If web articles are also provided below, you may additionally cite them (news, tooling, docs -- current/practical context, not peer-reviewed) via each section's cited_web_urls, but never in place of a paper citation -- keep the two kinds of sources clearly distinguished, the same way qa.py's chat answers do.

Use inline bracket markers in each section's content to mark which source supports each claim: [Paper 1], [Paper 2], ... for papers, in the order you list them in that section's own cited_paper_ids; [Web 1], [Web 2], ... for web articles (if any were provided), in the order you list them in that section's own cited_web_urls. These are two separate numbering sequences, never merged into one, and independent PER SECTION -- start over at [Paper 1]/[Web 1] in each of the three sections, never carried over from a previous section. These are temporary, section-local markers only -- they are automatically converted into one shared, report-wide [1], [2], [3]... numbering afterward, so you do not need to (and should not try to) coordinate numbers across sections yourself.
"""


# --- report-quality Phase R2B: Analytical dynamic section generation ---
# (defined here, ahead of _build_report_schema just below, since that
# function's own default parameter references _LEGACY_SECTION_DEFINITIONS)

# The exact same 3 sections SECTION_NAMES/SYSTEM_PROMPT above already
# describe, just repackaged as {key, title, description} definitions so
# _build_report_schema and the other section-iterating helpers below can
# be generalized to accept ANY section-definitions list while still
# defaulting to this one -- every existing caller/test that doesn't
# explicitly opt into Analytical sections keeps byte-identical behavior.
_LEGACY_SECTION_DEFINITIONS: list[dict] = [
    {"key": "findings", "title": "Findings", "description": FINDINGS_DESCRIPTION},
    {"key": "limitations", "title": "Limitations", "description": LIMITATIONS_DESCRIPTION},
    {"key": "future_scope", "title": "Future Scope", "description": FUTURE_SCOPE_DESCRIPTION},
]

# The approved 8-section Analytical report body -- word budgets are
# prompt GUIDANCE only (see _build_report_system_prompt's closing
# paragraph), never enforced by truncation.
REPORT_SECTION_DEFINITIONS: list[dict] = [
    {
        "key": "executive_summary", "title": "Executive Summary",
        "description": (
            "A concise, high-level summary of what this literature review covers and its most "
            "important takeaway(s) -- written so a reader could read only this section and still "
            "understand the gist. Target roughly 100-150 words."
        ),
    },
    {
        "key": "introduction_scope", "title": "Introduction & Scope",
        "description": (
            "Introduces the research topic and states plainly what this review covers -- exactly "
            "the hand-picked paper set given below, nothing broader. Target roughly 150-220 words."
        ),
    },
    {
        "key": "thematic_findings", "title": "Thematic Findings",
        "description": (
            "Synthesized findings and key contributions across the selected papers, grounded strictly "
            "in their abstracts. Organize by theme, not paper-by-paper, and synthesize across papers "
            "within each theme -- what do they collectively show, where do their emphases differ -- "
            "rather than summarizing each paper in isolation one after another. Target roughly "
            "350-550 words."
        ),
    },
    {
        "key": "methodology_landscape", "title": "Methodology Landscape",
        "description": (
            "Compares the methods/approaches the selected papers actually use -- similarities, "
            "differences, and tradeoffs between them, grounded strictly in their abstracts. Target "
            "roughly 180-300 words."
        ),
    },
    {
        "key": "contradictions_open_debates", "title": "Contradictions & Open Debates",
        "description": (
            "Disagreements, conflicting findings, methodological tensions, tradeoffs, or unresolved "
            "design choices across the selected papers -- direct contradictions are only one form "
            "this can take; papers more often differ by prioritizing different metrics, making "
            "different tradeoffs, or leaving a design question unresolved, and those count too. State "
            "explicitly if genuinely none of these are apparent across this specific paper set -- "
            "don't invent one. Target roughly 180-300 words."
        ),
    },
    {
        "key": "gap_analysis", "title": "Gap Analysis",
        "description": (
            "What the selected papers, taken together, do NOT cover or leave unaddressed -- missing "
            "populations, settings, comparisons, or evaluations, grounded in what's actually absent "
            "from their abstracts, not outside knowledge. Diagnostic, not prescriptive: identify and "
            "describe the absence itself, without proposing how to address it -- that belongs in "
            "Future Research Directions below, not here. Target roughly 180-300 words."
        ),
    },
    {
        "key": "future_research_directions", "title": "Future Research Directions",
        "description": (
            "Concrete future research directions plausibly implied by the selected papers -- may be "
            "more speculative than other sections, but must still stay grounded in what the papers "
            "actually cover, not outside knowledge. Prescriptive, not just diagnostic: propose what "
            "should actually be studied or built next, going beyond restating what Gap Analysis "
            "already identified as missing -- turn each relevant gap into an actionable direction, or "
            "propose a genuinely new angle the papers' own results suggest. Target roughly 150-250 "
            "words."
        ),
    },
    {
        "key": "conclusion", "title": "Conclusion",
        "description": (
            "A short closing synthesis -- not a repeat of the Executive Summary, a final word tying "
            "the review together. Target roughly 80-140 words."
        ),
    },
]

# report-quality Phase R2C: the Foundational template -- same 8 keys/
# titles as REPORT_SECTION_DEFINITIONS above (a deliberate, low-risk
# design constraint: no template ever introduces/renames/reorders a
# section, only varies its own instructions and word-budget guidance),
# written for a reader newer to the topic -- defines key concepts before
# using them, explains why a method/result matters, wider word-budget
# ranges where clarity genuinely needs the room. Still evidence-grounded
# and cited exactly like every other template (that requirement lives in
# the shared prompt scaffolding in _build_report_system_prompt, not
# here, so it can never be accidentally dropped per-template).
_FOUNDATIONAL_SECTION_DEFINITIONS: list[dict] = [
    {
        "key": "executive_summary", "title": "Executive Summary",
        "description": (
            "A concise, high-level summary of what this literature review covers and its most "
            "important takeaway(s), written so a reader new to the topic can follow it without prior "
            "background -- briefly say what the general area is about before summarizing the "
            "findings. Target roughly 120-180 words."
        ),
    },
    {
        "key": "introduction_scope", "title": "Introduction & Scope",
        "description": (
            "Introduces the research topic and states plainly what this review covers -- exactly the "
            "hand-picked paper set given below, nothing broader. Define any key terms or concepts a "
            "newcomer to this topic would need before the rest of the report makes sense. Target "
            "roughly 200-300 words."
        ),
    },
    {
        "key": "thematic_findings", "title": "Thematic Findings",
        "description": (
            "Synthesized findings and key contributions across the selected papers, grounded strictly "
            "in their abstracts. Organize by theme, not paper-by-paper, and synthesize across papers "
            "within each theme. For each theme, explain in plain terms why the finding or method "
            "matters, not just what it is. Target roughly 400-650 words."
        ),
    },
    {
        "key": "methodology_landscape", "title": "Methodology Landscape",
        "description": (
            "Compares the methods/approaches the selected papers actually use -- similarities, "
            "differences, and tradeoffs between them, grounded strictly in their abstracts. Briefly "
            "explain any method-specific terminology before comparing methods that use it. Target "
            "roughly 220-350 words."
        ),
    },
    {
        "key": "contradictions_open_debates", "title": "Contradictions & Open Debates",
        "description": (
            "Disagreements, conflicting findings, methodological tensions, tradeoffs, or unresolved "
            "design choices across the selected papers -- direct contradictions are only one form "
            "this can take; papers more often differ by prioritizing different metrics, making "
            "different tradeoffs, or leaving a design question unresolved, and those count too. "
            "Explain plainly why a given tension matters for someone learning this area. State "
            "explicitly if genuinely none of these are apparent across this specific paper set -- "
            "don't invent one. Target roughly 200-320 words."
        ),
    },
    {
        "key": "gap_analysis", "title": "Gap Analysis",
        "description": (
            "What the selected papers, taken together, do NOT cover or leave unaddressed -- missing "
            "populations, settings, comparisons, or evaluations, grounded in what's actually absent "
            "from their abstracts, not outside knowledge. Diagnostic, not prescriptive: identify and "
            "describe the absence itself, without proposing how to address it -- that belongs in "
            "Future Research Directions below, not here. Target roughly 200-320 words."
        ),
    },
    {
        "key": "future_research_directions", "title": "Future Research Directions",
        "description": (
            "Concrete future research directions plausibly implied by the selected papers -- may be "
            "more speculative than other sections, but must still stay grounded in what the papers "
            "actually cover, not outside knowledge. Prescriptive, not just diagnostic: propose what "
            "should actually be studied or built next, going beyond restating what Gap Analysis "
            "already identified as missing. Target roughly 170-280 words."
        ),
    },
    {
        "key": "conclusion", "title": "Conclusion",
        "description": (
            "A short closing synthesis -- not a repeat of the Executive Summary, a final word tying "
            "the review together for a reader who is now newly familiar with this topic. Target "
            "roughly 90-150 words."
        ),
    },
]

# report-quality Phase R2C: the Expert template -- same 8 keys/titles as
# REPORT_SECTION_DEFINITIONS above, written for a reader already
# confident in the topic -- denser, skips textbook background unless a
# specific comparison needs it, emphasizes cross-paper relationships,
# tensions, methodological nuance, unstated assumptions, and concrete
# research opportunities, tighter word-budget ranges than Foundational.
_EXPERT_SECTION_DEFINITIONS: list[dict] = [
    {
        "key": "executive_summary", "title": "Executive Summary",
        "description": (
            "A dense, high-level summary of what this literature review covers and its most "
            "important takeaway(s) for a reader already familiar with the area -- skip introductory "
            "framing, lead directly with the substance. Target roughly 80-130 words."
        ),
    },
    {
        "key": "introduction_scope", "title": "Introduction & Scope",
        "description": (
            "States plainly what this review covers -- exactly the hand-picked paper set given below, "
            "nothing broader -- without re-explaining background a reader already confident in this "
            "topic doesn't need. Target roughly 100-160 words."
        ),
    },
    {
        "key": "thematic_findings", "title": "Thematic Findings",
        "description": (
            "Synthesized findings and key contributions across the selected papers, grounded strictly "
            "in their abstracts. Organize by theme, not paper-by-paper, and synthesize across papers "
            "within each theme -- emphasize how findings relate to and qualify each other, not just "
            "what each one independently shows. Target roughly 320-500 words."
        ),
    },
    {
        "key": "methodology_landscape", "title": "Methodology Landscape",
        "description": (
            "Compares the methods/approaches the selected papers actually use, grounded strictly in "
            "their abstracts -- emphasize methodological nuance and the specific assumptions each "
            "approach makes, not just a surface-level similarities/differences list. Target roughly "
            "180-300 words."
        ),
    },
    {
        "key": "contradictions_open_debates", "title": "Contradictions & Open Debates",
        "description": (
            "Disagreements, conflicting findings, methodological tensions, tradeoffs, or unresolved "
            "design choices across the selected papers -- direct contradictions are only one form "
            "this can take; papers more often differ by prioritizing different metrics, making "
            "different tradeoffs, or leaving a design question unresolved, and those count too. "
            "Surface unstated assumptions behind a given position where the abstracts make them "
            "inferable. State explicitly if genuinely none of these are apparent across this specific "
            "paper set -- don't invent one. Target roughly 180-300 words."
        ),
    },
    {
        "key": "gap_analysis", "title": "Gap Analysis",
        "description": (
            "What the selected papers, taken together, do NOT cover or leave unaddressed -- missing "
            "populations, settings, comparisons, or evaluations, grounded in what's actually absent "
            "from their abstracts, not outside knowledge. Diagnostic, not prescriptive: identify and "
            "describe the absence itself, without proposing how to address it -- that belongs in "
            "Future Research Directions below, not here. Target roughly 160-260 words."
        ),
    },
    {
        "key": "future_research_directions", "title": "Future Research Directions",
        "description": (
            "Concrete, specific research opportunities plausibly implied by the selected papers -- "
            "may be more speculative than other sections, but must still stay grounded in what the "
            "papers actually cover, not outside knowledge. Prescriptive, not just diagnostic: propose "
            "what should actually be studied or built next, going beyond restating what Gap Analysis "
            "already identified as missing -- favor precise, actionable directions over broad "
            "statements. Target roughly 130-220 words."
        ),
    },
    {
        "key": "conclusion", "title": "Conclusion",
        "description": (
            "A short closing synthesis -- not a repeat of the Executive Summary, a final word tying "
            "the review together. Target roughly 70-120 words."
        ),
    },
]

# report-quality Phase R2C: the one place all three templates are
# collected -- REPORT_TEMPLATES[report_template] resolves a template
# name to its section-definitions list, the same list _build_report_
# schema/_build_report_system_prompt already accept generically. Kept as
# three independent, fully-spelled-out lists rather than one templated-
# with-overrides structure -- matches this module's own existing
# precedent (_LEGACY_SECTION_DEFINITIONS vs. REPORT_SECTION_DEFINITIONS
# were already two independent full lists, not one derived from the
# other) and keeps each template's actual prompt text trivially
# greppable/reviewable on its own.
REPORT_TEMPLATES: dict[str, list[dict]] = {
    "foundational": _FOUNDATIONAL_SECTION_DEFINITIONS,
    "analytical": REPORT_SECTION_DEFINITIONS,
    "expert": _EXPERT_SECTION_DEFINITIONS,
}

ANALYTICAL_SECTION_NAMES: tuple[str, ...] = tuple(d["key"] for d in REPORT_SECTION_DEFINITIONS)
# report-quality Phase R2C: despite the name (kept unchanged from R2B to
# minimize churn across existing call sites), this tuple of 8 keys is
# now shared by EVERY template, not just Analytical -- REPORT_TEMPLATES'
# three lists all use the exact same key set by design, so any one of
# them would produce an identical tuple here.

# Maps a legacy compatibility field to whichever NEW Analytical section
# it's projected from (see _project_legacy_fields) -- reused in reverse
# (_ANALYTICAL_KEY_TO_LEGACY_FIELD below) to resolve prior citations to
# preserve when regenerating a session whose EXISTING report still has
# the old 3-section shape under the new 8-section schema (see
# _resolve_prior_citations_for_regeneration). One decision, reused for
# both purposes, not two independently-chosen mappings that could
# disagree. The 5 Analytical sections with no entry here (executive_
# summary, introduction_scope, methodology_landscape, gap_analysis,
# conclusion) have no legacy analogue at all -- an old report never had
# anything corresponding to them.
_LEGACY_FIELD_TO_ANALYTICAL_KEY: dict[str, str] = {
    "findings": "thematic_findings",
    "limitations": "contradictions_open_debates",
    "future_scope": "future_research_directions",
}
_ANALYTICAL_KEY_TO_LEGACY_FIELD: dict[str, str] = {v: k for k, v in _LEGACY_FIELD_TO_ANALYTICAL_KEY.items()}


def _build_report_schema(
    paper_ids: list[str], web_urls: list[str] | None = None,
    section_definitions: list[dict] | None = None,
) -> type[BaseModel]:
    """Same technique as summarize.py's _build_response_schema/
    _build_web_response_schema (a per-call dynamic Literal restricted to
    the exact ids passed in) -- reimplemented here, not imported, since
    the schema SHAPE is genuinely different (named fixed sections vs. N
    theme clusters), not because the grounding mechanism differs.

    web_urls is optional-guarded exactly like qa.py's
    _build_answer_schema's own web_urls parameter -- a Literal can't be
    built from an empty tuple, and not every report has web sources, so
    cited_web_urls simply doesn't exist as a field unless there's
    something it could legitimately reference.

    section_definitions (report-quality Phase R2B) defaults to the
    original 3-section shape (_LEGACY_SECTION_DEFINITIONS) so every
    existing caller that doesn't pass it gets byte-identical behavior;
    generate_report/_regenerate_report_sections_with_sources explicitly
    pass REPORT_SECTION_DEFINITIONS instead. Named REQUIRED pydantic
    fields (one per section key) make a missing or duplicate section
    structurally impossible -- pydantic validation rejects it outright --
    which is why this stays a set of named fields rather than a
    Literal-constrained list of section items.
    """
    section_definitions = section_definitions if section_definitions is not None else _LEGACY_SECTION_DEFINITIONS
    paper_id_literal = Literal[tuple(paper_ids)]

    section_fields: dict = {
        "content": (str, Field(description="The section's write-up text")),
        "cited_paper_ids": (
            list[paper_id_literal],
            Field(description="Selected paper_ids that support this section's content, in citation order. May be empty if none specifically support it."),
        ),
    }
    if web_urls:
        web_url_literal = Literal[tuple(web_urls)]
        section_fields["cited_web_urls"] = (
            list[web_url_literal],
            Field(description="Web article urls that support this section's content, in citation order. May be empty."),
        )

    section_model = create_model("ReportSection", **section_fields)

    return create_model(
        "LiteratureReviewReport",
        **{d["key"]: (section_model, Field(description=d["description"])) for d in section_definitions},
    )


# report-quality Phase R2C: the ONE template-specific paragraph appended
# to the otherwise fully shared prompt scaffolding below -- this is
# deliberately the ONLY place templates diverge in the shared-scaffolding
# text (everything else -- grounding, citation density, web-source
# constraints, marker instructions -- stays identical across templates).
# "analytical" maps to "" (nothing appended) SPECIFICALLY so Analytical's
# generated prompt stays byte-identical to what R2B/R2B.1 already shipped
# -- see tests/test_report.py's own non-regression test for this.
_TEMPLATE_DEPTH_GUIDANCE: dict[str, str] = {
    "foundational": (
        "\n\nThis is a FOUNDATIONAL-depth report, written for a reader newer to this topic: define key "
        "concepts and background terms before using them, explain in plain terms why a method or "
        "result matters, and use simpler, more explicit transitions between ideas. Sections may run "
        "toward the upper end of their word budget (or slightly past it) where genuinely needed for "
        "clarity -- clarity takes priority over brevity here. Still ground and cite every claim exactly "
        "as instructed above; explanatory depth is not a license to skip evidence. Reminder: even when "
        "naming a paper's real identifier in your explanatory prose (e.g. \"the paper 2308.06821v1 "
        "shows...\"), the CITATION marker right after it must still be [Paper N] in the instructed "
        "format -- never repeat that identifier itself as the bracketed marker."
    ),
    "analytical": "",
    "expert": (
        "\n\nThis is an EXPERT-depth report, written for a reader already confident in this topic: skip "
        "introductory or textbook-level background explanation unless it's genuinely needed to make a "
        "specific comparison land, and prioritize depth over accessibility. Emphasize relationships "
        "between papers, tradeoffs, methodological nuance, unstated assumptions, and concrete research "
        "opportunities over restating what each paper does. Still ground and cite every claim exactly "
        "as instructed above; density is not a license to assert unsupported claims."
    ),
}


def _build_report_system_prompt(section_definitions: list[dict], template: str = DEFAULT_REPORT_TEMPLATE) -> str:
    """report-quality Phase R2B: builds the system prompt for an
    arbitrary section_definitions list -- generalizes the exact same
    framing/grounding/citation-marker instructions SYSTEM_PROMPT above
    hardcodes for the fixed 3-section case, over however many named
    sections are given. SYSTEM_PROMPT itself stays untouched (still
    _generate_report_sections' own default parameter value) -- this is a
    new, independent builder that generate_report/_regenerate_report_
    sections_with_sources now call explicitly with REPORT_SECTION_
    DEFINITIONS, not a refactor of SYSTEM_PROMPT's own text.

    report-quality Phase R2C: `template` selects one paragraph from
    _TEMPLATE_DEPTH_GUIDANCE, appended after the shared scaffolding below
    -- everything else in this function's returned text is identical
    across templates (only the passed-in section_definitions' own
    per-section description/word-budget text differs by template,
    resolved by the caller via REPORT_TEMPLATES before this is called).
    """
    numbered_sections = "\n".join(
        f"{i}. {d['title']}: {d['description']}" for i, d in enumerate(section_definitions, start=1)
    )
    plural = "s" if len(section_definitions) != 1 else ""
    return f"""You are writing a literature review report over a specific, deliberately hand-picked set of papers for a research topic. You will be given the topic and exactly this set of papers (id, title, abstract only) -- this is the FINAL set a user has already chosen; do not second-guess or add papers beyond it.

Write exactly {len(section_definitions)} section{plural}, in this order:
{numbered_sections}

For EACH section, ground every claim STRICTLY in the given abstracts -- do not use outside knowledge about these papers, their authors, or the topic beyond what the abstracts state. List which paper_ids actually support that section's content in its own cited_paper_ids, in the order you'd reference them. A section's citations are independent of every other section's -- a paper cited in one section does not need to also appear in another, and vice versa. Some sections may end up citing few or no papers directly if they're mostly synthesizing across the review as a whole (e.g. an executive summary or conclusion) -- that's fine, don't force citations that aren't genuinely there.

Cite specific, evidence-bearing claims as they occur, not just once per paragraph or only for a paper's overall thesis: whenever you name a specific method, technique, dataset, metric, result, or direct comparison drawn from a paper, mark it inline right there. Do not, however, force a marker onto purely transitional or synthesis sentences that make no new, specific claim -- citation density should track where the evidence actually is, not appear on every sentence mechanically.

If web articles are also provided below, you may additionally cite them (news, tooling, docs -- current/practical context, not peer-reviewed) via each section's cited_web_urls, but never in place of a paper citation, and never as your primary source of evidence for a section -- the selected papers are always the primary evidence base. Only cite a web article when it directly supports or extends a claim already grounded in the papers (e.g. a tool's current real-world adoption, a technique's practical deployment status, a recent development building on a paper's method) -- do not cite a web article merely because it is topically adjacent, and do not use one to introduce a new topic the paper set itself doesn't cover. If none of the given web articles are directly relevant to a section's content, cite none there -- an empty cited_web_urls for a section is expected and fine.

Use inline bracket markers in each section's content to mark which source supports each claim: [Paper 1], [Paper 2], ... for papers, in the order you list them in that section's own cited_paper_ids; [Web 1], [Web 2], ... for web articles (if any were provided), in the order you list them in that section's own cited_web_urls. These are two separate numbering sequences, never merged into one, and independent PER SECTION -- start over at [Paper 1]/[Web 1] in each section, never carried over from a previous section. These are temporary, section-local markers only -- they are automatically converted into one shared, report-wide [1], [2], [3]... numbering afterward, so you do not need to (and should not try to) coordinate numbers across sections yourself. When a single claim is directly supported by more than one source (e.g. comparing two papers' methods in the same sentence), mark each one with its own bracket back-to-back, e.g. [Paper 2][Paper 5] -- never combine multiple citations into one bracket like [Paper 2, Paper 5]. Never cite a source using its paper_id, DOI, arXiv ID, Semantic Scholar ID, URL, or title inside the bracket -- always [Paper N] or [Web N] exactly as instructed above, even if you've just named that source's real identifier or title in your own prose right before the marker. The conversion into final numeric citations only recognizes this exact [Paper N]/[Web N] format; anything else in a bracket is not a citation the system can resolve.

Word budgets given above are guidance, not a hard limit -- stay roughly within them, erring toward being concise and evidence-dense rather than padding to hit a count, especially if the given paper set is small. Avoid generic, textbook-style phrasing that could apply to any paper set on this general topic -- ground each claim in what THESE specific selected papers actually say, not a general summary of the field.
""" + _TEMPLATE_DEPTH_GUIDANCE.get(template, "")


# --- report-quality Phase R1: inline numbered citations + References ---

# Matches a section-local, temporary marker as prompted above -- e.g.
# "[Paper 2]" or "[Web 1]" -- AND (report-quality Phase R2C citation-
# marker fix) a GROUPED marker bundling two or more comma-separated
# entries in one bracket, e.g. "[Paper 6, Paper 8]" or "[Paper 3, Web
# 1]": the model is instructed to write one marker per citation, but in
# practice sometimes bundles several into a single bracket when a claim
# is backed by more than one source (observed most often in sections
# that explicitly ask for cross-paper comparison, like Methodology
# Landscape). A bracket that doesn't match this shape at all (anything
# other than one-or-more "Paper N"/"Web N" entries separated by commas)
# is left untouched, same as before. Deliberately reimplemented here
# rather than imported from qa.py's own _CITATION_MARKER_RE: the
# single-marker pattern itself is identical, but this module already
# reimplements _build_report_schema for the same reason (see that
# function's own docstring) -- avoiding a coupling on qa.py's private
# internals for what's a genuinely self-contained regex.
_SECTION_CITATION_MARKER_RE = re.compile(
    r"\[\s*(?:Paper|Web)\s+\d+(?:\s*,\s*(?:Paper|Web)\s+\d+)*\s*\]"
)
# Extracts each individual (kind, number) entry out of a marker matched
# by _SECTION_CITATION_MARKER_RE above -- run via .findall() on the
# whole matched bracket text, so it works identically whether that
# bracket held one entry or several.
_SECTION_CITATION_MARKER_ENTRY_RE = re.compile(r"(Paper|Web)\s+(\d+)")

# report-quality Phase R2C citation-hardening fix: a deterministic
# backstop for the model ignoring the instructed [Paper N]/[Web N]
# format and citing a source by its own real identifier instead --
# observed specifically in Foundational-template output, e.g.
# "[2308.06821v1]" or "[abd1c342495432171beb7ca8fd9551ef13cbd0ff]"
# (a paper's arXiv id / Semantic Scholar-style hash id) leaking straight
# into the rendered report instead of resolving to a numbered citation.
# Matches a single bracket whose content has no whitespace (a real
# paper_id/DOI/URL is always one unbroken token; ordinary bracketed
# prose almost always contains a space, which keeps this from
# misfiring on some unrelated aside the model happens to bracket) and
# that does NOT already look like a valid "Paper N"/"Web N" marker (the
# negative lookahead below) -- that shape is left alone for
# _SECTION_CITATION_MARKER_RE/_densify_section_markers to handle exactly
# as before. See _resolve_raw_source_id_markers below for the actual
# exact-match-against-known-sources resolution.
_RAW_SOURCE_ID_MARKER_RE = re.compile(r"\[(?!\s*(?:Paper|Web)\s+\d+)([^\s\[\]]+)\]")


def _resolve_raw_source_id_markers(
    content: str, section_name: str, cited_papers: list[Paper], cited_web_articles: list[WebArticle],
    assigner: _ReferenceAssigner, section_marked_keys: dict[str, set[tuple[str, str]]],
) -> str:
    """report-quality Phase R2C citation-hardening fix: resolves a raw-
    identifier marker (see _RAW_SOURCE_ID_MARKER_RE) directly to the
    SAME global reference number the source would get through the
    normal [Paper N]/[Web N] path -- both go through the same shared
    `assigner`, so a source cited once via its raw id and again via a
    correct [Paper N] marker elsewhere in the report still collapses to
    exactly one reference number, not two.

    A bracket only ever resolves here if its trimmed content EXACTLY
    matches one of THIS section's own cited_papers' paper_id or
    cited_web_articles' url -- never a fuzzy/partial/global-pool match,
    the same "no guessing" discipline _build_references_and_renumber's
    own out-of-range handling already uses. A bare digit string (the
    shape of an already-final "[N]" marker) is deliberately never
    treated as a candidate id -- this app's real paper_ids/urls are
    never plain digits, and it's cheap insurance against ever
    misinterpreting an unrelated numeric bracket.

    Runs BEFORE _densify_section_markers/the regular marker-resolve pass,
    directly on the section's ORIGINAL, unmodified content -- resolving
    straight to a final "[N]" string, which neither of those two
    passes' own regexes matches, so it's left untouched by them
    afterward. A raw id that matches nothing known is INVALID, same
    "strip it, don't guess" policy as an out-of-range [Paper N] marker:
    dropped entirely (never left as raw text in the final report body),
    logged as a warning.
    """
    paper_by_id = {p.paper_id: p for p in cited_papers}
    article_by_url = {a.url: a for a in cited_web_articles}

    def _replace(match: re.Match[str]) -> str:
        candidate = match.group(1)
        if candidate.isdigit():
            return match.group()
        if candidate in paper_by_id:
            paper = paper_by_id[candidate]
            section_marked_keys[section_name].add(("paper", paper.paper_id))
            number = assigner.get_or_assign("paper", paper.paper_id, paper.title, paper=paper)
            return f"[{number}]"
        if candidate in article_by_url:
            article = article_by_url[candidate]
            section_marked_keys[section_name].add(("web", article.url))
            number = assigner.get_or_assign("web", article.url, article.title, article=article)
            return f"[{number}]"
        logger.warning(
            "_resolve_raw_source_id_markers: dropping unrecognized bracketed marker %r in %r section "
            "(matches no cited paper_id or web url there)", candidate, section_name,
        )
        return ""

    return _RAW_SOURCE_ID_MARKER_RE.sub(_replace, content)


# report-quality Phase R2B.1: an invalid/out-of-range marker resolves to
# "" in _build_references_and_renumber's own _resolve (see that
# function's docstring) -- the space that used to separate the marker
# from surrounding text is left behind, e.g. "classification [Paper 9]."
# becomes "classification ." once the bracket is stripped. These two
# regexes clean that up; deliberately narrow (collapse repeated spaces,
# drop a space immediately before punctuation) rather than a general
# prose reformatter.
_MULTI_SPACE_RE = re.compile(r" {2,}")
_SPACE_BEFORE_PUNCTUATION_RE = re.compile(r" +([,.;:!?])")


def _cleanup_marker_stripped_whitespace(content: str) -> str:
    """report-quality Phase R2B.1: applied only to freshly generated
    report content, inside _build_references_and_renumber, AFTER marker
    resolution/stripping -- never to derive_legacy_references or any
    other old-report compatibility path, which promise never to rewrite
    an old report's prose at all. Idempotent and a no-op on already-
    clean text (no markers were ever stripped), so it's safe to run
    unconditionally on every section rather than only when a marker was
    actually dropped."""
    content = _MULTI_SPACE_RE.sub(" ", content)
    return _SPACE_BEFORE_PUNCTUATION_RE.sub(r"\1", content)


def _densify_section_markers(content: str) -> str:
    """Renumbers ONE section's own [Paper N]/[Web N] markers to be dense
    and sequential starting at 1, per kind, in the order they first
    appear in the text -- same defensive technique qa.py's own
    _renumber_citation_markers uses, and for the same reason: the
    model's raw marker numbering isn't guaranteed to already align with
    its cited_paper_ids/cited_web_urls list order (observed skipping
    numbers in practice there; not assumed reliable here either). This
    is a prerequisite step for _build_references_and_renumber below --
    it does NOT resolve markers to actual paper_ids/urls or produce the
    final report-wide numbering, only guarantees each section's own
    markers are clean (1, 2, 3, ... with no gaps) before that mapping is
    attempted.

    report-quality Phase R2C citation-marker fix: a GROUPED raw marker
    like "[Paper 6, Paper 8]" is split into separate single-entry
    brackets here ("[Paper 2][Paper 3]", densified) rather than kept
    grouped -- every marker downstream of this function (including
    _build_references_and_renumber's own re-use of
    _SECTION_CITATION_MARKER_RE) then only ever needs to reason about
    however many entries a given bracket happens to hold, single or
    multiple, with no special-casing required.
    """
    next_number = {"Paper": 1, "Web": 1}
    seen: dict[tuple[str, str], int] = {}

    def _replace(match: re.Match[str]) -> str:
        parts = []
        for kind, old_number in _SECTION_CITATION_MARKER_ENTRY_RE.findall(match.group()):
            key = (kind, old_number)
            if key not in seen:
                seen[key] = next_number[kind]
                next_number[kind] += 1
            parts.append(f"[{kind} {seen[key]}]")
        return "".join(parts)

    return _SECTION_CITATION_MARKER_RE.sub(_replace, content)


class _ReferenceAssigner:
    """Assigns (or reuses) a report-wide reference number for a given
    source, appending a new ReferenceEntry-shaped dict to `references`
    the first time each source is seen, and exposing the (kind, key) ->
    number map so a caller can look up an already-assigned number
    without re-deriving it. Shared by _build_references_and_renumber and
    derive_legacy_references below so the actual reference-entry-
    building logic (citation formatting, link_url resolution) lives in
    exactly one place, not duplicated between the "fresh generation" and
    "old persisted report" paths.
    """

    def __init__(self, references: list[dict]) -> None:
        self.references = references
        self.number_by_key: dict[tuple[str, str], int] = {}

    def get_or_assign(self, kind: str, key: str, title: str, *, paper: Paper | None = None, article: WebArticle | None = None) -> int:
        existing = self.number_by_key.get((kind, key))
        if existing is not None:
            return existing
        number = len(self.references) + 1
        self.number_by_key[(kind, key)] = number
        if kind == "paper":
            assert paper is not None
            link_url = f"https://doi.org/{paper.doi}" if paper.doi else paper.url
            self.references.append({
                "number": number, "kind": "paper", "paper_id": key, "url": None,
                "title": title, "formatted": format_apa_citation(paper), "link_url": link_url,
            })
        else:
            assert article is not None
            self.references.append({
                "number": number, "kind": "web", "paper_id": None, "url": key,
                "title": title, "formatted": format_web_citation(article), "link_url": key,
            })
        return number


def _build_references_and_renumber(sections_out: dict, section_names: tuple[str, ...] = SECTION_NAMES) -> dict:
    """The R1 post-processing pass: converts each section's temporary,
    section-local [Paper N]/[Web N] markers into ONE global, bare-number
    [N] sequence shared across the whole report, and builds the
    top-level `references` list those numbers point to. Deterministic,
    no LLM call -- operates purely on the already-attached
    cited_papers/cited_web_articles per section (real Paper/WebArticle
    objects, already in citation order), so it's directly testable
    without going through model-calling machinery at all.

    MUST run after _restore_dropped_citations on the regeneration path
    (i.e. after cited_papers/cited_web_articles are already final) --
    this function only reads citations and rewrites content/reference_
    numbers, it never changes WHICH papers/web articles are cited, so
    it's safe to call exactly once, as the very last step, regardless of
    caller.

    section_names (report-quality Phase R2B) defaults to the original
    3-tuple SECTION_NAMES; generate_report/_regenerate_report_sections_
    with_sources pass ANALYTICAL_SECTION_NAMES instead so the same
    algorithm below runs unchanged over all 8 Analytical sections.

    Algorithm:
      1. Per section, in the given section_names order: densify that
         section's own markers (see
         _densify_section_markers), then resolve each densified marker
         to the underlying paper_id/url via that section's own
         cited_papers[n-1]/cited_web_articles[n-1] (0-indexed from the
         1-based densified marker number), assigning/reusing a GLOBAL
         reference number for it. The first time a source is seen
         anywhere in the report, it gets the next global number; every
         later occurrence (same section or a different one) reuses that
         same number -- "the same source keeps the same number across
         the whole report."
      2. A densified marker number beyond that section's own cited list
         length (e.g. [Paper 3] when only 2 papers were cited there) is
         INVALID -- stripped from the text entirely (never guessed at),
         and logged.
      3. Only AFTER every section's markers are resolved (so marker-
         derived citations always get the lowest/earliest numbers): any
         source a section structurally cited but never actually marked
         inline still gets a trailing global number and References
         entry, appended in that section's own citation-list order --
         a citation the model forgot to bracket is never silently
         dropped from References.
    """
    references: list[dict] = []
    assigner = _ReferenceAssigner(references)
    section_marked_keys: dict[str, set[tuple[str, str]]] = {name: set() for name in section_names}
    rewritten_content: dict[str, str] = {}

    for section_name in section_names:
        section = sections_out[section_name]
        cited_papers = section["cited_papers"]
        cited_web_articles = section.get("cited_web_articles", [])
        # report-quality Phase R2C citation-hardening fix: resolves any
        # raw-identifier marker (see _resolve_raw_source_id_markers'
        # own docstring) BEFORE densify/the regular [Paper N]/[Web N]
        # pass -- straight to a final "[N]" string, which neither of
        # those two passes' own regexes matches, so it passes through
        # them unchanged afterward.
        raw_id_resolved = _resolve_raw_source_id_markers(
            section["content"], section_name, cited_papers, cited_web_articles, assigner, section_marked_keys,
        )
        densified = _densify_section_markers(raw_id_resolved)

        def _resolve(
            match: re.Match[str], section_name: str = section_name,
            cited_papers: list[Paper] = cited_papers, cited_web_articles: list[WebArticle] = cited_web_articles,
        ) -> str:
            # report-quality Phase R2C citation-marker fix: densify above
            # already split every raw marker into single-entry brackets,
            # but this function is written to handle 1+ entries per
            # bracket regardless -- resolving each entry independently
            # and re-joining as separate "[N]" brackets (never "[N, M]",
            # see _SECTION_CITATION_MARKER_ENTRY_RE's own docstring for
            # why: the frontend's own marker regex only ever recognizes
            # single-number brackets). An entry whose index is out of
            # range is dropped on its own, not the whole bracket -- e.g.
            # "[Paper 1, Paper 9]" with only one paper cited there
            # resolves to just "[<n>]", the same "strip what's invalid,
            # keep what's valid" policy the old single-marker code used,
            # just applied per-entry instead of per-bracket.
            parts = []
            for kind, densified_number_str in _SECTION_CITATION_MARKER_ENTRY_RE.findall(match.group()):
                densified_number = int(densified_number_str)
                index = densified_number - 1
                if kind == "Paper":
                    if index >= len(cited_papers):
                        logger.warning(
                            "_build_references_and_renumber: dropping out-of-range marker [Paper %d] in %r section "
                            "(only %d paper(s) cited there)", densified_number, section_name, len(cited_papers),
                        )
                        continue
                    paper = cited_papers[index]
                    section_marked_keys[section_name].add(("paper", paper.paper_id))
                    number = assigner.get_or_assign("paper", paper.paper_id, paper.title, paper=paper)
                else:
                    if index >= len(cited_web_articles):
                        logger.warning(
                            "_build_references_and_renumber: dropping out-of-range marker [Web %d] in %r section "
                            "(only %d web article(s) cited there)", densified_number, section_name, len(cited_web_articles),
                        )
                        continue
                    article = cited_web_articles[index]
                    section_marked_keys[section_name].add(("web", article.url))
                    number = assigner.get_or_assign("web", article.url, article.title, article=article)
                parts.append(f"[{number}]")
            return "".join(parts)

        rewritten_content[section_name] = _SECTION_CITATION_MARKER_RE.sub(_resolve, densified)

    # Trailing pass: structurally cited but never marked -- still counts.
    for section_name in section_names:
        section = sections_out[section_name]
        for paper in section["cited_papers"]:
            if ("paper", paper.paper_id) not in section_marked_keys[section_name]:
                assigner.get_or_assign("paper", paper.paper_id, paper.title, paper=paper)
        for article in section.get("cited_web_articles", []):
            if ("web", article.url) not in section_marked_keys[section_name]:
                assigner.get_or_assign("web", article.url, article.title, article=article)

    for section_name in section_names:
        section = sections_out[section_name]
        section["content"] = _cleanup_marker_stripped_whitespace(rewritten_content[section_name])
        all_keys = (
            section_marked_keys[section_name]
            | {("paper", p.paper_id) for p in section["cited_papers"]}
            | {("web", a.url) for a in section.get("cited_web_articles", [])}
        )
        section["reference_numbers"] = sorted(assigner.number_by_key[key] for key in all_keys)

    return {**sections_out, "references": references}


def build_references_and_renumber(sections_out: dict, section_names: tuple[str, ...] = SECTION_NAMES) -> dict:
    """report-quality Phase R3.2 Chunk 2: a public, stable name for
    _build_references_and_renumber above, for reuse by callers outside
    this module -- curation_chat.py's derive_chat_references being the
    first (chat needs the exact same grouped-marker/raw-source-id/
    whitespace-cleanup/global-numbering behavior, over "sections" that
    are really per-exchange chat turns, not report sections). A pure
    pass-through, not a reimplementation -- _build_references_and_
    renumber itself is intentionally left in place, unrenamed, so every
    existing report.py call site and test keeps working unchanged;
    this is purely an additive, alternate entry point.

    Each call still builds its own brand-new `references` list and
    `_ReferenceAssigner` internally (see that function's own docstring)
    -- there is no shared/global numbering state between a caller using
    this wrapper and report.py's own report-generation call sites, so a
    chat-scoped and a report-scoped call can never see or influence each
    other's numbering, regardless of how many times either runs.
    """
    return _build_references_and_renumber(sections_out, section_names)


def derive_legacy_references(report: dict) -> dict:
    """report-quality Phase R1 backward compatibility: a report persisted
    before this phase has no `references`/`reference_numbers` at all
    (see curation_session.py's _deserialize_report) and its `content`
    was never generated with any inline markers to rewrite. Called by
    the API serializer (research_agent/api_app/serializers.py's
    _report_to_out), not the storage layer, so old persisted dicts stay
    an untouched, byte-identical round-trip and this derivation is
    recomputed fresh on every read instead of being silently migrated
    into storage.

    Same numbering policy as _build_references_and_renumber's own
    trailing "structurally cited but unmarked" pass (Pass 3 there),
    applied to EVERY citation in an old report, since none of them were
    ever inline-marked in the first place: global numbers assigned in
    section order (findings, limitations, future_scope), then citation-
    list order within each section. `content` is returned unchanged --
    an old report's prose is never retroactively rewritten with markers
    it was never generated with.
    """
    references: list[dict] = []
    assigner = _ReferenceAssigner(references)

    updated_sections = {}
    for section_name in SECTION_NAMES:
        section = report[section_name]
        numbers: list[int] = []
        for paper in section["cited_papers"]:
            numbers.append(assigner.get_or_assign("paper", paper.paper_id, paper.title, paper=paper))
        for article in section.get("cited_web_articles", []):
            numbers.append(assigner.get_or_assign("web", article.url, article.title, article=article))
        updated_sections[section_name] = {**section, "reference_numbers": sorted(set(numbers))}

    return {**report, **updated_sections, "references": references}


# --- report-quality Phase R2A: dynamic section model (schema/rendering
# migration only -- generation itself is unchanged, still exactly the
# fixed findings/limitations/future_scope pipeline above) ---

_LEGACY_SECTION_TITLES = {"findings": "Findings", "limitations": "Limitations", "future_scope": "Future Scope"}


def derive_sections_from_legacy_report(report: dict) -> list[dict]:
    """Every report today -- old persisted ones, and every currently-
    generated one, since generation hasn't changed yet -- only ever has
    the fixed findings/limitations/future_scope shape. This derives the
    new `sections` list (ReportOut.sections, the future primary report
    body) straight from that existing shape, 1:1 and in order. Called by
    the API serializer (api_app/serializers.py's _report_to_out)
    whenever a report dict has no "sections" key of its own yet -- in
    practice, every report right now. Same "derive at read time, never
    rewrite persisted storage" philosophy as derive_legacy_references
    above -- never mutates `report`.

    Once real multi-section generation ships (a later phase), a fresh
    report will carry its own genuine `sections` list and this function
    simply won't be reached for it (see _report_to_out's `"sections" not
    in report` guard) -- this is a bridge for the pre-R2B/pre-R2C world,
    not a replacement for real dynamic generation.
    """
    return [
        {
            "key": name,
            "title": _LEGACY_SECTION_TITLES[name],
            "content": report[name]["content"],
            "reference_numbers": report[name].get("reference_numbers", []),
        }
        for name in SECTION_NAMES
    ]


@observe(name="generate_report_sections", as_type="generation", capture_input=False, capture_output=False)
def _generate_report_sections(
    topic: str, papers: list[Paper], web_articles: list[WebArticle], schema: type[BaseModel],
    client: OpenAI, model: str = REPORT_MODEL, system_prompt: str = SYSTEM_PROMPT,
) -> BaseModel:
    paper_listing = "\n\n".join(
        f"paper_id: {p.paper_id}\ntitle: {p.title}\nabstract: {p.abstract or '(no abstract available)'}"
        for p in papers
    )
    context_sections = [f"Selected papers:\n\n{paper_listing}"]
    if web_articles:
        web_listing = "\n\n".join(
            f"url: {a.url}\ntitle: {a.title}\nsnippet: {a.snippet or '(no snippet available)'}"
            for a in web_articles
        )
        context_sections.append(f"Additional web sources:\n\n{web_listing}")

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Research topic: {topic}\n\n" + "\n\n".join(context_sections)},
    ]

    langfuse = get_client()
    langfuse.update_current_generation(input=messages, model=model)

    response = client.chat.completions.parse(model=model, messages=messages, response_format=schema)
    parsed = response.choices[0].message.parsed
    if parsed is None:
        langfuse.update_current_generation(output=None, level="WARNING", status_message="Model refused")
        raise RuntimeError(f"Model refused to produce a report: {response.choices[0].message.refusal}")

    usage = response.usage
    if usage is not None:
        logger.info(
            "generate_report: %d tokens billed (prompt=%d, completion=%d)",
            usage.total_tokens, usage.prompt_tokens, usage.completion_tokens,
        )
        langfuse.update_current_generation(
            usage_details={
                "input_tokens": usage.prompt_tokens,
                "output_tokens": usage.completion_tokens,
                "total_tokens": usage.total_tokens,
            },
        )
    langfuse.update_current_generation(output=parsed.model_dump())
    return parsed


def _project_legacy_fields(sections_out: dict) -> dict:
    """report-quality Phase R2B: stamps the three legacy findings/
    limitations/future_scope keys onto a freshly-Analytical-generated
    sections_out dict, as straight aliases of their mapped Analytical
    section (see _LEGACY_FIELD_TO_ANALYTICAL_KEY) -- not independently
    generated text of their own. Exists purely so callers that still
    read report["findings"] directly (ReportOut's own required findings/
    limitations/future_scope fields, api_app/serializers.py's
    _report_to_out) keep working unchanged for a genuinely new report,
    without the model ever being asked to write three additional,
    redundant sections. Must run AFTER _build_references_and_renumber --
    the projected dict needs each section's FINAL content/
    reference_numbers, not the pre-renumbering ones.
    """
    return {
        **sections_out,
        **{
            legacy_key: dict(sections_out[analytical_key])
            for legacy_key, analytical_key in _LEGACY_FIELD_TO_ANALYTICAL_KEY.items()
        },
    }


def _sections_list(sections_out: dict) -> list[dict]:
    """Builds the ReportOut.sections-shaped list (the real, generated
    counterpart to R2A's derive_sections_from_legacy_report) straight
    from a finalized sections_out dict -- one entry per
    REPORT_SECTION_DEFINITIONS key, in order, carrying each section's
    already-renumbered content/reference_numbers."""
    return [
        {
            "key": d["key"], "title": d["title"],
            "content": sections_out[d["key"]]["content"],
            "reference_numbers": sections_out[d["key"]]["reference_numbers"],
        }
        for d in REPORT_SECTION_DEFINITIONS
    ]


@observe(name="generate_report", capture_input=False, capture_output=False)
def generate_report(
    topic: str, selected_papers: list[Paper], web_articles: list[WebArticle] | None = None,
    client: OpenAI | None = None, model: str = REPORT_MODEL,
    report_template: str = DEFAULT_REPORT_TEMPLATE,
) -> dict:
    """report-quality Phase R2B: generates the real 8-section report
    body over EXACTLY selected_papers -- the model is structurally
    unable to cite any paper_id outside this set (Literal-constrained,
    same guarantee level as summarize.py's own citations, and as this
    report's own pre-R2B 3-section body). Returns one dict entry per
    section key ({"content", "cited_papers", ...}), a top-level
    "sections" list (the real ReportOut.sections the API serializer
    already reads), the legacy findings/limitations/future_scope keys as
    straight PROJECTIONS of their mapped section (see
    _project_legacy_fields) -- not independently generated -- plus
    "skipped_papers" (selected papers never cited in ANY section,
    mirroring summarize.py's own convention, logged as a warning not an
    error), "references" (report-quality Phase R1's global numbered
    citations, now spanning all 8 sections), and "report_template" (the
    resolved template name, report-quality Phase R2C -- the report dict
    itself is the one source of truth for which template produced it).

    report_template (report-quality Phase R2C) selects which of
    REPORT_TEMPLATES' three section-definitions lists governs this
    call's prompt/schema -- defaults to DEFAULT_REPORT_TEMPLATE
    ("analytical"), so every existing caller that doesn't pass it keeps
    generating exactly what it always has.

    web_articles (curation-chat-web-escalation Phase 5d) is optional and
    backward-compatible: every existing caller that doesn't pass it gets
    byte-identical behavior to before -- each section dict simply has no
    "cited_web_articles" key at all unless web_articles is non-empty,
    the same convention _build_report_schema() uses for the schema
    field itself.
    """
    section_definitions = REPORT_TEMPLATES[report_template]
    web_articles = web_articles or []
    if not selected_papers:
        empty_section = {"content": "", "cited_papers": [], "reference_numbers": []}
        if web_articles:
            empty_section["cited_web_articles"] = []
        sections_out = {d["key"]: dict(empty_section) for d in section_definitions}
        result = _project_legacy_fields(sections_out)
        result["sections"] = _sections_list(sections_out)
        result["skipped_papers"] = []
        result["references"] = []
        result["report_template"] = report_template
        return result

    papers_by_id = {p.paper_id: p for p in selected_papers}
    web_by_url = {a.url: a for a in web_articles}
    schema = _build_report_schema(list(papers_by_id), list(web_by_url) or None, section_definitions)
    system_prompt = _build_report_system_prompt(section_definitions, report_template)

    client = client or OpenAI()
    parsed = _generate_report_sections(topic, selected_papers, web_articles, schema, client, model, system_prompt=system_prompt)

    referenced_ids: set[str] = set()
    sections_out = {}
    for section_name in ANALYTICAL_SECTION_NAMES:
        section = getattr(parsed, section_name)
        cited_papers = [papers_by_id[pid] for pid in section.cited_paper_ids]  # guaranteed present: Literal enforced it structurally
        referenced_ids.update(section.cited_paper_ids)
        section_out = {"content": section.content, "cited_papers": cited_papers}
        if web_by_url:
            cited_web_urls = list(getattr(section, "cited_web_urls", []))
            section_out["cited_web_articles"] = [web_by_url[url] for url in cited_web_urls]
        sections_out[section_name] = section_out

    skipped = [p for pid, p in papers_by_id.items() if pid not in referenced_ids]
    if skipped:
        logger.warning(
            "generate_report: %d selected paper(s) never cited in any section: %s",
            len(skipped), [p.title for p in skipped],
        )

    get_client().update_current_span(
        input={"topic": topic, "num_selected_papers": len(selected_papers), "num_web_articles": len(web_articles)},
        output={
            **{f"{name}_citation_count": len(sections_out[name]["cited_papers"]) for name in ANALYTICAL_SECTION_NAMES},
            "skipped_papers": paper_metadata(skipped),
        },
    )

    result = _build_references_and_renumber({**sections_out, "skipped_papers": skipped}, ANALYTICAL_SECTION_NAMES)
    result = _project_legacy_fields(result)
    result["sections"] = _sections_list(result)
    result["report_template"] = report_template
    return result


def generate_report_for_session(
    session: PaperPoolSession, client: OpenAI | None = None, model: str = REPORT_MODEL,
    report_template: str = DEFAULT_REPORT_TEMPLATE,
) -> dict:
    """Session-aware wrapper: refuses cleanly if the session isn't
    actually ready for synthesis yet (stage != "synthesize") rather than
    generating prematurely over a still-in-progress pick set, then
    resolves selected_papers directly from session.selected_papers
    (already full Paper data, populated at pick-time in
    curation_loop.py -- NOT reconstructed from session.reserve, which
    may no longer contain already-served papers after a refill).
    """
    if session.stage != "synthesize":
        raise ValueError(
            f"Session is not ready for report synthesis (stage={session.stage!r}, expected 'synthesize') -- "
            "curation must finish (target met, user stopped, or topic exhausted) before generating a report."
        )
    return generate_report(session.topic, session.selected_papers, client=client, model=model, report_template=report_template)


def _resolve_prior_citations_for_regeneration(existing_report: dict, section_key: str) -> list[Paper]:
    """report-quality Phase R2B: resolves which papers (if any) were
    cited under `section_key` in the PRIOR report, for citation-
    preservation purposes. Handles both the ordinary case (existing_
    report already has this exact key -- e.g. a second Analytical
    regeneration) and the cross-version case (existing_report still has
    the OLD 3-section shape, e.g. its first report predates R2B):
    thematic_findings/contradictions_open_debates/future_research_
    directions fall back to their mapped legacy key's prior citations
    (_ANALYTICAL_KEY_TO_LEGACY_FIELD) if the new key itself isn't
    present. Every other new section key has no old-report analogue at
    all -- an old report never had anything to preserve for e.g.
    Executive Summary or Gap Analysis -- so those simply resolve to an
    empty prior-citations list, which _restore_dropped_citations already
    treats as a valid, unremarkable case (nothing to restore)."""
    if section_key in existing_report:
        return existing_report[section_key]["cited_papers"]
    legacy_field = _ANALYTICAL_KEY_TO_LEGACY_FIELD.get(section_key)
    if legacy_field is not None and legacy_field in existing_report:
        return existing_report[legacy_field]["cited_papers"]
    return []


def _build_regeneration_system_prompt(
    existing_report: dict, section_definitions: list[dict], template: str = DEFAULT_REPORT_TEMPLATE,
    allowed_web_urls: set[str] | None = None, web_by_url: dict[str, WebArticle] | None = None,
) -> str:
    """report-quality Phase R3.1b: allowed_web_urls/web_by_url are
    optional -- when given, appends a paragraph naming the user-approved
    web sources by title and instructing the model to cite one inline
    with [Web N] if (and only if) it's genuinely relevant, or omit it
    entirely otherwise. This is now the ONLY mechanism that gives an
    approved web source special treatment during regeneration -- there
    is no deterministic backstop that adds one to the report if the
    model doesn't cite it inline itself (see _regenerate_report_
    sections_with_sources' own docstring for why: a metadata-only
    "restored"/"force-included" citation with no inline marker reads as
    a broken/decorative reference to an actual reader, so R3.1b removed
    that guarantee rather than keep producing it)."""
    per_section_citations = "\n".join(
        f"- {d['key']}: {[p.paper_id for p in _resolve_prior_citations_for_regeneration(existing_report, d['key'])]}"
        for d in section_definitions
    )
    approved_web_paragraph = ""
    if allowed_web_urls and web_by_url:
        approved_titles = [web_by_url[url].title for url in allowed_web_urls if url in web_by_url]
        if approved_titles:
            titles_list = "\n".join(f"- {title}" for title in approved_titles)
            approved_web_paragraph = f"""

The following web source(s) were specifically approved by the user via chat, not merely discovered by search:
{titles_list}
Integrate one of these only where it is directly relevant to a specific claim you are already making in a section -- if you use it, you MUST cite it inline with a [Web N] marker exactly as instructed above, at the claim it actually supports. If an approved source is not directly relevant to what a section is about, omit it entirely rather than citing it just because it was approved. There is no separate mechanism that adds an approved source to the report on your behalf -- if you don't cite it inline, it will not appear anywhere in the report."""
    return _build_report_system_prompt(section_definitions, template) + f"""
This is a REGENERATION of an existing report -- additional web sources have been approved since it was first written, and you should incorporate them where genuinely relevant. You MUST preserve every paper_id already cited in a section below: never drop an existing paper citation from a section just because new web sources are now available. You may add new web citations, cite additional papers, and refine the prose, but an already-cited paper must remain cited in any section that previously cited it.

Previously cited paper_ids per section:
{per_section_citations}
""" + approved_web_paragraph


def _restore_dropped_citations(existing_report: dict, section_name: str, cited_paper_ids: list[str]) -> list[str]:
    """Layer 2 of the two-layer citation-preservation guarantee (same
    two-layer pattern as Phase 1c: prompt instruction first, structural
    enforcement as the backstop that doesn't depend on the model
    actually following it). Appends back, in their original order, any
    paper_id this exact section cited in the PRIOR report but the
    regeneration's own output dropped -- restoring is unconditional, not
    a best-effort nudge, so it holds even if the prompt instruction
    above is ignored entirely (see tests/test_report.py's isolation test,
    which defeats the prompt layer on purpose to prove this layer alone
    is what's actually load-bearing).

    report-quality Phase R2B: prior citations are resolved via
    _resolve_prior_citations_for_regeneration rather than direct
    existing_report[section_name] access, so this also works correctly
    (never KeyErrors) when existing_report still has the old 3-section
    shape and section_name is a new Analytical key."""
    restored = list(cited_paper_ids)
    for pid in [p.paper_id for p in _resolve_prior_citations_for_regeneration(existing_report, section_name)]:
        if pid not in restored:
            restored.append(pid)
    return restored


def _regenerate_report_sections_with_sources(
    session: PaperPoolSession, web_articles: list[WebArticle], client: OpenAI | None, model: str, caller_name: str,
    report_template: str | None = None, allowed_web_urls: set[str] | None = None,
) -> dict:
    """Shared body for regenerate_report_with_new_sources (whole-pool) and
    regenerate_report_with_approved_web_sources (curation-chat-add-to-
    report Phase 4's selective subset) -- identical schema-building,
    citation-restoration, and skipped-paper logic either way; the only
    difference between the two public callers is WHICH web_articles list
    gets passed in here. Neither public function duplicates this body, so
    citation handling can never drift between the two paths. caller_name
    is only for the skipped-papers log line, so it stays attributable.

    Same preconditions as before this refactor (session.report must
    already exist; stage must be "synthesize") -- unchanged, just moved
    here from regenerate_report_with_new_sources's own body.

    allowed_web_urls (report-quality Phase R3.1b, replacing R2D/R3.1's
    force-include/restore guarantee): the set of web URLs to flag to the
    model, in the regeneration prompt, as user-approved via chat (see
    _build_regeneration_system_prompt's own approved_web_paragraph).
    This is now PURELY a prompt signal -- there is no deterministic
    post-processing that adds or restores a web citation the model
    doesn't itself put in a section's own cited_web_urls this round.
    R2D/R3.1 originally guaranteed an approved source's presence
    regardless of whether the model cited it, via _restore_dropped_web_
    citations (metadata-only restoration of a prior citation) and
    _force_include_allowed_web_articles (metadata-only injection of a
    never-cited one) -- both were removed in R3.1b because a citation
    with cited_web_articles metadata but no inline [N] marker in that
    section's own content reads to a reader as a broken/decorative
    reference (an orphan), exactly the reported bug. The CURRENT
    round's own model output is now the sole source of truth for which
    web sources a section cites: this function's own cited_web_urls
    handling below just passes the model's own selection through
    (filtered to non-revoked candidates, same as before). Deliberately
    a SEPARATE set from `web_articles` (the candidates the model is
    merely OFFERED) -- `web_articles` stays whatever it always was for
    each caller (the full raw pool for regenerate_report_with_new_
    sources, exactly approved_web_articles for regenerate_report_with_
    approved_web_sources). None/omitted (the default) just means no
    approved-sources paragraph is added to the prompt. Already revoked-
    safe with zero extra filtering needed here -- a revoked URL is
    excluded from `web_articles` (and therefore from web_by_url, and
    therefore from both the prompt paragraph and anything the model
    could possibly cite) by the existing revocation filter below.

    report_template (report-quality Phase R2C): None (the default) means
    PRESERVE session.report's own current template, resolved below via
    existing_report.get("report_template", DEFAULT_REPORT_TEMPLATE) --
    an old report predating this phase has no "report_template" key at
    all, correctly defaulting to Analytical. An explicit value switches
    the report to that template for this regeneration. This same
    None-means-preserve default is what makes curation_chat_service.py's
    add-to-report flow (which never passes report_template at all)
    automatically preserve the existing template with zero changes
    needed at that call site.

    report-quality Phase R2D citation-revocation fix: before building
    the schema/prompt, any web_articles entry whose URL is in session.
    revoked_web_article_urls is dropped entirely -- structurally, not
    just probabilistically, so the model is never even offered a
    revoked source as a citable candidate. session.revoked_web_article_
    urls (curation_chat.py's delete_chat_exchanges/edit_chat_exchange
    populate it; _accept_web_offer clears an entry the moment it's
    re-cited) is the PERSISTENT record of "this URL lost its only live
    chat backing" -- deliberately not re-derived here from session.
    report's own current references, which would self-heal to "nothing
    to revoke" after just one clean regeneration and let a later
    regeneration re-offer (and the model re-cite) the same source, since
    session.web_articles_added itself is never pruned. Applies to BOTH
    regenerate_report_with_new_sources and regenerate_report_with_
    approved_web_sources uniformly -- for the latter this is normally a
    no-op (approved_web_articles is already filtered against revocation
    upstream by curation_chat.py's resolve_approved_web_articles_for_
    regeneration), just a defensive backstop, not a behavior change.
    """
    if session.report is None:
        raise ValueError(
            "Session has no existing report to regenerate -- call generate_report_for_session() first."
        )
    if session.stage != "synthesize":
        raise ValueError(
            f"Session is not ready for report synthesis (stage={session.stage!r}, expected 'synthesize') -- "
            "curation must finish (target met, user stopped, or topic exhausted) before generating a report."
        )

    existing_report = session.report
    selected_papers = session.selected_papers
    resolved_template = report_template if report_template is not None else existing_report.get("report_template", DEFAULT_REPORT_TEMPLATE)
    section_definitions = REPORT_TEMPLATES[resolved_template]
    allowed_web_urls = allowed_web_urls or set()

    if session.revoked_web_article_urls:
        web_articles = [a for a in web_articles if a.url not in session.revoked_web_article_urls]

    papers_by_id = {p.paper_id: p for p in selected_papers}
    web_by_url = {a.url: a for a in web_articles}
    schema = _build_report_schema(list(papers_by_id), list(web_by_url) or None, section_definitions)
    system_prompt = _build_regeneration_system_prompt(
        existing_report, section_definitions, resolved_template,
        allowed_web_urls=allowed_web_urls, web_by_url=web_by_url,
    )

    client = client or OpenAI()
    parsed = _generate_report_sections(
        session.topic, selected_papers, web_articles, schema, client, model, system_prompt=system_prompt,
    )

    referenced_ids: set[str] = set()
    sections_out = {}
    for section_name in ANALYTICAL_SECTION_NAMES:
        section = getattr(parsed, section_name)
        cited_paper_ids = _restore_dropped_citations(existing_report, section_name, list(section.cited_paper_ids))

        referenced_ids.update(cited_paper_ids)
        cited_papers = [papers_by_id[pid] for pid in cited_paper_ids]
        section_out = {"content": section.content, "cited_papers": cited_papers}
        if web_by_url:
            # report-quality Phase R3.1b: the model's own cited_web_urls
            # this round is the sole source of truth -- no restoration,
            # no force-inclusion. A url not in the model's own current
            # output simply isn't cited this round, same as any other
            # candidate it chose not to use.
            cited_web_urls = list(getattr(section, "cited_web_urls", []))
            section_out["cited_web_articles"] = [web_by_url[url] for url in cited_web_urls if url in web_by_url]
        sections_out[section_name] = section_out

    skipped = [p for pid, p in papers_by_id.items() if pid not in referenced_ids]
    if skipped:
        logger.warning(
            "%s: %d selected paper(s) never cited in any section: %s",
            caller_name, len(skipped), [p.title for p in skipped],
        )

    get_client().update_current_span(
        input={
            "topic": session.topic, "num_selected_papers": len(selected_papers),
            "num_web_articles": len(web_articles),
        },
        output={
            **{f"{name}_citation_count": len(sections_out[name]["cited_papers"]) for name in ANALYTICAL_SECTION_NAMES},
            "skipped_papers": paper_metadata(skipped),
        },
    )

    # report-quality Phase R1: renumbering/reference-building runs LAST,
    # after _restore_dropped_citations above has already finalized which
    # papers each section cites -- a restored PAPER citation the model's
    # own regenerated content never bracketed correctly falls into this
    # pass's own "structurally cited but never marked" handling rather
    # than being silently missing from References. Web citations have no
    # such restoration step as of R3.1b (see _regenerate_report_
    # sections_with_sources' own docstring) -- cited_web_articles above
    # already reflects only what the model itself cited this round.
    result = _build_references_and_renumber({**sections_out, "skipped_papers": skipped}, ANALYTICAL_SECTION_NAMES)
    result = _project_legacy_fields(result)
    result["sections"] = _sections_list(result)
    result["report_template"] = resolved_template
    return result


def regenerate_report_with_new_sources(
    session: PaperPoolSession, client: OpenAI | None = None, model: str = REPORT_MODEL,
    report_template: str | None = None,
) -> dict:
    """curation-chat-web-escalation Phase 5d: regenerates session.report
    over the SAME session.selected_papers plus ALL of
    session.web_articles_added approved so far via Phase 5c's
    offer-and-decide mechanism -- typically called after a new web
    source is approved into an already-synthesized session, not a
    from-scratch generation (that's generate_report_for_session()).

    Requires session.report to already exist -- refuses cleanly
    otherwise, since "regenerate" implies something to regenerate FROM;
    generate_report_for_session() is the right call for a session's
    first report.

    Never mutates session.report itself -- returns the new report dict,
    same convention as generate_report()/generate_report_for_session();
    the caller decides when to actually replace session.report with it.

    curation-chat-add-to-report Phase 4: this function's behavior is
    UNCHANGED by that phase -- still whole-pool, still what POST
    /curation/{id}/report/regenerate uses. The selective, approved-
    subset-only path added in Phase 4 is the separate
    regenerate_report_with_approved_web_sources() below; the two are
    intentionally independent (see that function's own docstring for
    why using both on the same session has a real interaction worth
    knowing about).

    report-quality Phase R2D citation-revocation fix: still whole-pool
    over session.web_articles_added -- but any URL in session.revoked_
    web_article_urls (curation_chat.py's delete_chat_exchanges/edit_
    chat_exchange populate it whenever a URL loses its only live chat
    backing) is excluded from the candidates offered this call, even
    though it's still sitting in the raw, unfiltered web_articles_added
    pool -- see _regenerate_report_sections_with_sources's own
    docstring for why this needs to be a persistent, session-level
    record rather than something re-derived from the prior report alone.
    A source that was simply never revoked is unaffected -- "whole pool"
    still means every web article chat has ever surfaced, just no longer
    including one whose only chat backing has since been deleted or
    edited away.

    report-quality Phase R2C: report_template=None (the default)
    preserves session.report's own current template; an explicit value
    switches the report to that template for this regeneration. See
    _regenerate_report_sections_with_sources's own docstring for the
    exact resolution rule.

    report-quality Phase R3.1b (previously R2D/R3.1's web-citation-
    preservation fix): allowed_web_urls is session.report_approved_web_
    article_urls, NOT the full web_articles_added pool passed as
    candidates above -- the model may still cite ANY web article ever
    discovered in chat, same as always, but only a source that's been
    through the explicit "Add to report" approval flow is named in the
    regeneration prompt's approved-sources paragraph (see
    _build_regeneration_system_prompt), encouraging the model to cite
    it inline if relevant. There is no deterministic guarantee anymore
    -- an approved source the model doesn't cite this round simply
    doesn't appear, same as a merely-discovered-but-never-approved one
    always has. This is a deliberate change from R2D/R3.1: their
    force-include/restore guarantee produced a References entry with no
    inline marker for a never-cited source, which read as a broken/
    decorative reference to an actual reader -- R3.1b removed that
    guarantee rather than keep producing it.
    """
    return _regenerate_report_sections_with_sources(
        session, session.web_articles_added, client, model, "regenerate_report_with_new_sources",
        report_template=report_template, allowed_web_urls=set(session.report_approved_web_article_urls),
    )


def regenerate_report_with_approved_web_sources(
    session: PaperPoolSession, approved_web_articles: list[WebArticle],
    client: OpenAI | None = None, model: str = REPORT_MODEL,
    report_template: str | None = None,
) -> dict:
    """curation-chat-add-to-report Phase 4: regenerates session.report over
    session.selected_papers plus ONLY approved_web_articles -- deliberately
    does NOT read session.web_articles_added (the raw, unfiltered pool of
    every web article ever discovered during chat) at all. The caller
    (curation_chat.py's resolve_approved_web_articles_for_regeneration)
    is responsible for filtering the raw pool down to the approved subset
    BEFORE calling this; an unapproved web article can never reach the
    model through this function no matter what else is sitting in the
    session's raw pool.

    Same preconditions/return convention as regenerate_report_with_new_
    sources (session.report must already exist, stage must be
    "synthesize", never mutates session.report itself).

    Interaction worth knowing: if a session ever also uses the OLD
    whole-pool regenerate_report_with_new_sources (e.g. via the Report
    tab's existing "Regenerate" button), that call will overwrite
    session.report with one reflecting the ENTIRE raw pool, including any
    web articles never approved through this selective path -- the two
    mechanisms don't defer to each other. Not resolved in Phase 4 (kept
    as an explicit, named follow-up), since /report/regenerate's own
    behavior must stay whole-pool and unchanged this phase. (Partially
    addressed by report-quality Phase R2D: whole-pool regeneration no
    longer resurrects a web citation whose only chat backing has since
    been deleted/edited away -- see regenerate_report_with_new_sources's
    own docstring. The two mechanisms are still independent in every
    other respect: a whole-pool regeneration can still include a web
    article never approved through THIS selective path at all.)

    report-quality Phase R2D citation-revocation fix: the shared
    regeneration body's session.revoked_web_article_urls exclusion also
    applies here, though it's normally a no-op in practice for this
    function specifically -- curation_chat.py's resolve_approved_web_
    articles_for_regeneration already keeps approved_web_articles itself
    correct against Phase B pruning before it ever reaches this call, so
    a revoked URL should never even be IN approved_web_articles to begin
    with. A defensive backstop, not a behavior change for this function.

    report-quality Phase R2C: report_template exists on this function's
    signature purely for symmetry with regenerate_report_with_new_
    sources and because the shared body needs it either way -- the ONE
    real caller (curation_chat.py's add-to-report flow, via
    curation_chat_service.py's add_curation_chat_exchanges_to_report)
    never passes it, so this always resolves to "preserve the existing
    report's current template" (the default, report_template=None). The
    add-to-report HTTP endpoint itself has no report_template field on
    its own request schema at all -- chat-triggered regeneration is not
    a product moment for choosing a template.

    report-quality Phase R3.1b (previously R2D/R3.1's web-citation-
    preservation fix): allowed_web_urls is exactly {a.url for a in
    approved_web_articles} -- every web article this call offers the
    model IS, by this function's own contract, already approved, so the
    entire candidate set is named in the regeneration prompt's approved-
    sources paragraph. As of R3.1b, an approved source the model doesn't
    genuinely cite inline this round does NOT appear in References --
    see _regenerate_report_sections_with_sources' own docstring for why
    the earlier force-include guarantee was removed (it produced an
    orphan reference: present in References, no inline marker anywhere
    in the body, which reads as broken/decorative to a reader).
    """
    return _regenerate_report_sections_with_sources(
        session, approved_web_articles, client, model, "regenerate_report_with_approved_web_sources",
        report_template=report_template, allowed_web_urls={a.url for a in approved_web_articles},
    )


# --- report-quality Phase R4.1: optional, bounded draft -> evaluate ->
# revise -> finalize refinement loop ---
#
# Implemented as plain, bounded functions -- NOT a LangGraph StateGraph.
# This flow has exactly one conditional fork (revise or don't) and no
# cycle: a revision, if it happens, is never re-evaluated, so there is
# nothing here that loops. LangGraph earns its keep elsewhere in this
# codebase specifically for checkpointing across multiple separate HTTP
# requests (curation_loop.py's pick loop) or interrupt()-based pausing
# for a real human response (planned, not yet built, for a future R4.4)
# -- neither applies to a flow that runs entirely, synchronously, inside
# one request. A bounded multi-round loop (R4.3) is still just a plain
# `for` loop when it lands; a StateGraph conversion is reserved for
# R4.4, if/when interrupt() is actually needed, mirroring how qa.py
# itself stayed plain functions until its own checkpointing/human-in-
# the-loop need became concrete, not before.
#
# Reuses every citation/reference/grounding helper generate_report and
# _regenerate_report_sections_with_sources already established --
# _build_report_schema (the dynamic Literal that makes citing outside
# the offered paper/web set structurally impossible, not just
# discouraged), _generate_report_sections, _restore_dropped_citations
# (paper preservation), and the full build_references_and_renumber ->
# _project_legacy_fields -> _sections_list post-processing chain.
# Nothing about citation/reference handling is reimplemented here.

_UNRESOLVED_MARKER_RE = re.compile(r"\[[^\[\]]*[A-Za-z][^\[\]]*\]")


class ReportEvaluation(BaseModel):
    """report-quality Phase R4.1: the evaluator's structured output.
    report.py-internal -- not an API schema (api_app/schemas.py has its
    own, much smaller ReportRefinementOut for what actually gets
    exposed; see evaluate_report's own docstring for why the two stay
    separate)."""

    overall_score: int = Field(
        description="Overall quality score from 0 (unusable) to 100 (excellent), judging synthesis "
        "quality, grounding, and appropriateness for the given template."
    )
    needs_revision: bool = Field(description="Whether this report should be revised before being finalized.")
    issues: list[str] = Field(
        default_factory=list,
        description="Specific, concrete issues found -- one per item, not a single vague paragraph.",
    )
    revision_instructions: str = Field(
        default="",
        description="Concrete, actionable instructions for how to fix the issues above. Empty string if "
        "needs_revision is false.",
    )
    section_scores: dict[str, int] | None = Field(
        default=None, description="Optional per-section quality score (0-100), keyed by section key.",
    )


_EVALUATOR_SYSTEM_PROMPT = """You are evaluating a literature review report for quality, before it is finalized and shown to a user. You will be given the report's topic, its reader-depth template (foundational/analytical/expert), the selected papers and web sources it was grounded in, and the report's own rendered sections.

Judge the report against these dimensions:
- Citation coverage: are claims actually backed by inline citations, not asserted bare?
- Source grounding: do claims genuinely reflect what the cited papers/web sources say, not outside knowledge?
- Section completeness: does each section actually deliver what its own description promises?
- Synthesis vs. paper-by-paper summary: does the report synthesize across sources by theme, or just summarize one paper at a time?
- Gap Analysis vs. Future Research Directions: are these two sections meaningfully distinct, not overlapping restatements of each other?
- Template appropriateness: does the depth/tone match the stated template (foundational = accessible/explanatory, expert = dense/assumes familiarity)?
- Web-source relevance: are cited web sources genuinely relevant to the specific claim they're attached to, not merely topically adjacent?
- Readability: is the prose clear and well-organized?

You will also be given a list of issues already found by a separate, deterministic check (structural problems like missing sections or broken citation markers) -- these are already known and do not need to be re-discovered, but should inform your own overall_score and needs_revision judgment.

Be specific in `issues` -- name the section and the actual problem, not vague generalities like "could be better." If the report is genuinely solid, say so and set needs_revision to false; do not manufacture issues to seem thorough.
"""


def _deterministic_report_checks(report: dict, report_template: str) -> tuple[list[str], list[str]]:
    """report-quality Phase R4.1: structural checks the evaluator's LLM
    judgment doesn't get a vote on -- objectively checkable, pure
    Python, no LLM call. Returns (hard_issues, warning_issues).

    hard_issues ALWAYS force needs_revision=True in evaluate_report --
    "do not rely only on vague LLM critique" for anything objectively
    checkable. warning_issues are informational only: fed to the LLM
    evaluator as context and included in the final issues list, but
    never force revision on their own (skipped_papers specifically --
    see below -- is expected, documented behavior in plenty of cases,
    not automatically a defect).

    In normal operation these mostly act as regression tripwires, not
    expected failure modes: raw grounding is already structurally
    guaranteed by _build_report_schema's dynamic Literal, and orphan
    web references are already prevented by construction as of R3.1b
    (see that section's own docstring) -- these checks exist to catch
    a FUTURE regression of either guarantee, not because either is
    expected to fail today.
    """
    section_definitions = REPORT_TEMPLATES[report_template]
    hard_issues: list[str] = []

    missing = [d["key"] for d in section_definitions if not report.get(d["key"], {}).get("content", "").strip()]
    if missing:
        hard_issues.append(f"Missing or empty required section(s): {missing}")

    leaked: set[str] = set()
    for d in section_definitions:
        content = report.get(d["key"], {}).get("content", "")
        leaked.update(_UNRESOLVED_MARKER_RE.findall(content))
    if leaked:
        hard_issues.append(f"Unresolved/raw citation marker(s) leaked into report body: {sorted(leaked)}")

    references = report.get("references", [])
    referenced_numbers = {n for d in section_definitions for n in report.get(d["key"], {}).get("reference_numbers", [])}
    orphans = [r["number"] for r in references if r["number"] not in referenced_numbers]
    if orphans:
        hard_issues.append(f"Orphan reference(s) in References with no citing section: {orphans}")

    numbers = sorted(r["number"] for r in references)
    if numbers and numbers != list(range(1, len(numbers) + 1)):
        hard_issues.append(f"Reference numbering is not a clean 1..{len(numbers)} sequence: {numbers}")

    warning_issues: list[str] = []
    skipped_papers = report.get("skipped_papers", [])
    if skipped_papers:
        warning_issues.append(
            f"{len(skipped_papers)} selected paper(s) never cited in any section: "
            f"{[p.title for p in skipped_papers]}"
        )

    return hard_issues, warning_issues


def _build_evaluation_prompt(
    report: dict, topic: str, selected_papers: list[Paper], web_articles: list[WebArticle],
    report_template: str, known_issues: list[str],
) -> str:
    section_definitions = REPORT_TEMPLATES[report_template]
    sections_text = "\n\n".join(
        f"[{d['key']}] {d['title']}\n{report.get(d['key'], {}).get('content', '')}"
        for d in section_definitions
    )
    paper_listing = "\n\n".join(
        f"paper_id: {p.paper_id}\ntitle: {p.title}\nabstract: {p.abstract or '(no abstract available)'}"
        for p in selected_papers
    )
    web_listing = "\n\n".join(
        f"url: {a.url}\ntitle: {a.title}\nsnippet: {a.snippet or '(no snippet available)'}"
        for a in web_articles
    ) or "(none)"
    known_issues_text = "\n".join(f"- {issue}" for issue in known_issues) or "(none)"

    return f"""Topic: {topic}
Template: {report_template}

Selected papers:

{paper_listing}

Web sources:

{web_listing}

Report sections:

{sections_text}

Issues already found by deterministic checks:
{known_issues_text}
"""


def _evaluate_report_llm(
    report: dict, topic: str, selected_papers: list[Paper], web_articles: list[WebArticle],
    report_template: str, known_issues: list[str], client: OpenAI, model: str = REPORT_MODEL,
) -> ReportEvaluation:
    messages = [
        {"role": "system", "content": _EVALUATOR_SYSTEM_PROMPT},
        {"role": "user", "content": _build_evaluation_prompt(
            report, topic, selected_papers, web_articles, report_template, known_issues,
        )},
    ]

    langfuse = get_client()
    langfuse.update_current_generation(input=messages, model=model)

    response = client.chat.completions.parse(model=model, messages=messages, response_format=ReportEvaluation)
    parsed = response.choices[0].message.parsed
    if parsed is None:
        langfuse.update_current_generation(output=None, level="WARNING", status_message="Model refused")
        raise RuntimeError(f"Model refused to evaluate the report: {response.choices[0].message.refusal}")
    langfuse.update_current_generation(output=parsed.model_dump())
    return parsed


@observe(name="evaluate_report", capture_input=False, capture_output=False)
def evaluate_report(
    report: dict, topic: str, selected_papers: list[Paper], web_articles: list[WebArticle],
    report_template: str, client: OpenAI, model: str = REPORT_MODEL,
) -> dict:
    """report-quality Phase R4.1: combines deterministic structural
    checks (hard gates) with an LLM judgment call (synthesis quality,
    grounding, template fit, readability -- genuinely subjective
    dimensions a regex can't assess). Returns a plain dict shaped like
    ReportEvaluation, not the pydantic model itself -- matches this
    module's own established "unpack to a plain dict past the raw
    OpenAI parse step" convention (see generate_report's own handling
    of its parsed response).

    Deterministic hard_issues are ALWAYS appended to `issues` and
    ALWAYS force needs_revision=True, regardless of what the LLM
    itself concluded -- objectively checkable structural problems are
    not up for the model's opinion. warning_issues (currently just
    skipped_papers) are appended to `issues` and passed to the LLM as
    context, but never force revision by themselves -- a section
    legitimately citing few/no papers is expected, documented behavior,
    not automatically a defect.
    """
    hard_issues, warning_issues = _deterministic_report_checks(report, report_template)
    known_issues = hard_issues + warning_issues
    evaluation = _evaluate_report_llm(
        report, topic, selected_papers, web_articles, report_template, known_issues, client, model,
    )
    result = evaluation.model_dump()
    result["issues"] = known_issues + result["issues"]
    if hard_issues:
        result["needs_revision"] = True
    return result


def _build_revision_system_prompt(
    draft: dict, evaluation: dict, section_definitions: list[dict], template: str = DEFAULT_REPORT_TEMPLATE,
) -> str:
    per_section_citations = "\n".join(
        f"- {d['key']}: {[p.paper_id for p in draft[d['key']]['cited_papers']]}"
        for d in section_definitions
    )
    draft_sections = "\n\n".join(
        f"[{d['key']}]\n{draft[d['key']]['content']}"
        for d in section_definitions
    )
    issues_text = "\n".join(f"- {issue}" for issue in evaluation["issues"]) or "(none)"

    return _build_report_system_prompt(section_definitions, template) + f"""
This is a REVISION of an existing draft, produced over the EXACT SAME selected papers and web sources offered below -- do not introduce new sources, and you structurally cannot cite anything outside what's offered. You MUST preserve every paper_id already cited in a section below: never drop an existing paper citation from a section just because you're revising it.

Previously cited paper_ids per section:
{per_section_citations}

The current draft, section by section:
{draft_sections}

An evaluator reviewed this draft and found the following issues:
{issues_text}

Revision instructions: {evaluation["revision_instructions"] or "(none given -- use your own judgment to address the issues above)"}

Rewrite the sections to address these issues while preserving what already works. Follow the exact same citation marker conventions described above.
"""


@observe(name="revise_report", capture_input=False, capture_output=False)
def revise_report(
    topic: str, selected_papers: list[Paper], web_articles: list[WebArticle],
    draft: dict, evaluation: dict, client: OpenAI, model: str = REPORT_MODEL,
    report_template: str = DEFAULT_REPORT_TEMPLATE,
) -> dict:
    """report-quality Phase R4.1: revises `draft` -- a report dict of
    the exact same shape generate_report/_regenerate_report_sections_
    with_sources produce -- over the EXACT SAME selected_papers/
    web_articles the draft was built from. Deliberately no new source
    discovery: the schema built below is constrained to precisely this
    paper/web set, the same dynamic-Literal grounding guarantee every
    other generation path already relies on, so a revision structurally
    cannot cite outside what the draft itself could cite.

    Paper citation preservation reuses _restore_dropped_citations
    exactly as regeneration does, with `draft` playing the same
    "existing_report" role a regeneration's own prior report plays --
    a paper cited in the draft stays cited in the revision even if the
    model's own revised output drops it.

    report-quality Phase R3.1b: web citations follow the identical
    product rule regeneration already established -- THIS round's own
    cited_web_urls is the sole source of truth, no restoration of a web
    citation the revision happens to drop, so a revision can never
    reintroduce an orphan References entry (nor can it force one in --
    there is no force-include mechanism here, same as regeneration
    since R3.1b).

    Runs through the exact same build_references_and_renumber ->
    _project_legacy_fields -> _sections_list post-processing chain as
    every other generation path -- the output is an ordinary report
    dict, not a distinct "revision" shape; a caller (or a later
    evaluate_report call) can't tell a revised report apart from a
    freshly generated one just by its structure.
    """
    section_definitions = REPORT_TEMPLATES[report_template]
    papers_by_id = {p.paper_id: p for p in selected_papers}
    web_by_url = {a.url: a for a in web_articles}
    schema = _build_report_schema(list(papers_by_id), list(web_by_url) or None, section_definitions)
    system_prompt = _build_revision_system_prompt(draft, evaluation, section_definitions, report_template)

    parsed = _generate_report_sections(
        topic, selected_papers, web_articles, schema, client, model, system_prompt=system_prompt,
    )

    referenced_ids: set[str] = set()
    sections_out = {}
    for section_name in ANALYTICAL_SECTION_NAMES:
        section = getattr(parsed, section_name)
        cited_paper_ids = _restore_dropped_citations(draft, section_name, list(section.cited_paper_ids))
        referenced_ids.update(cited_paper_ids)
        cited_papers = [papers_by_id[pid] for pid in cited_paper_ids if pid in papers_by_id]
        section_out = {"content": section.content, "cited_papers": cited_papers}
        if web_by_url:
            cited_web_urls = list(getattr(section, "cited_web_urls", []))
            section_out["cited_web_articles"] = [web_by_url[url] for url in cited_web_urls if url in web_by_url]
        sections_out[section_name] = section_out

    skipped = [p for pid, p in papers_by_id.items() if pid not in referenced_ids]
    if skipped:
        logger.warning(
            "revise_report: %d selected paper(s) never cited in any section: %s",
            len(skipped), [p.title for p in skipped],
        )

    result = _build_references_and_renumber({**sections_out, "skipped_papers": skipped}, ANALYTICAL_SECTION_NAMES)
    result = _project_legacy_fields(result)
    result["sections"] = _sections_list(result)
    result["report_template"] = report_template
    return result


def refine_report_if_requested(
    draft: dict, topic: str, selected_papers: list[Paper], web_articles: list[WebArticle],
    report_template: str, refinement_mode: str | None, client: OpenAI, model: str = REPORT_MODEL,
) -> dict:
    """report-quality Phase R4.1: the one orchestration point every
    report-generating call site that wants refinement goes through --
    mirrors append_report_version's own "one shared point" role for
    versioning below. Callers pass `draft` (whatever generate_report_
    for_session/regenerate_report_with_new_sources already produced)
    and the exact same selected_papers/web_articles/report_template
    that draft was built from.

    refinement_mode=None/"off" (DEFAULT_REFINEMENT_MODE) is a pure
    passthrough: `draft` is returned completely UNCHANGED -- not even
    a "refinement" key added -- with zero extra LLM calls. This is what
    keeps every existing caller's behavior byte-identical.

    refinement_mode="single": draft -> evaluate -> (revise once, if
    evaluate_report says needs_revision) -> finalize. Bounded by
    construction, not by a counter that could be miscounted: there is
    exactly one call to revise_report in this function's entire body,
    reached from exactly one branch, with no loop back to evaluate_
    report afterward -- an evaluator that would still flag issues after
    the one revision is never consulted again, the revision is
    finalized regardless. See this section's own module-level comment
    for why this is implemented as plain functions, not a LangGraph
    StateGraph.

    Stamps a `refinement` dict onto the returned report:
    {"enabled": True, "rounds": 0 or 1, "initial_score": int,
    "final_score": int | None, "issues": list[str],
    "revision_instructions": str, "section_scores": dict[str, int] |
    None}. `final_score` is None whenever a revision happened -- R4.1
    deliberately does not re-evaluate the revision, so there is no real
    score to report for it; inventing one would be worse than admitting
    it's unknown. `final_score` equals `initial_score` when no revision
    was needed (the draft IS the final report).

    report-quality Phase R4.2: `issues`/`revision_instructions`/
    `section_scores` come straight from the ONE evaluation that ever
    runs -- there is no second evaluation after a revision (see above),
    so these ALWAYS describe the DRAFT, not necessarily the finalized
    report a reader is looking at. If a revision happened, they're "what
    prompted the fix," not "what's still wrong" -- callers (the API
    serializer, the frontend) must not present them as a live critique
    of the current content. `section_scores` is None whenever the LLM
    evaluator didn't return one (ReportEvaluation's own field is
    optional); an empty/partial dict is also possible and must be
    tolerated, not assumed to cover every section key.
    """
    if refinement_mode != "single":
        return draft

    evaluation = evaluate_report(draft, topic, selected_papers, web_articles, report_template, client, model)
    initial_score = evaluation["overall_score"]

    if not evaluation["needs_revision"]:
        final = draft
        rounds = 0
        final_score = initial_score
    else:
        final = revise_report(
            topic, selected_papers, web_articles, draft, evaluation, client, model, report_template,
        )
        rounds = 1
        final_score = None

    final["refinement"] = {
        "enabled": True, "rounds": rounds, "initial_score": initial_score, "final_score": final_score,
        "issues": evaluation["issues"], "revision_instructions": evaluation["revision_instructions"],
        "section_scores": evaluation["section_scores"],
    }
    return final


# --- report-quality Phase R3: report versioning ---
#
# Every function above this section is unchanged in shape/behavior --
# generate_report/_regenerate_report_sections_with_sources still just
# return a fresh report dict, no session-mutation awareness of their
# own. append_report_version below is the ONE shared point every call
# site that used to do `session.report = <one of the functions above>`
# now goes through instead -- see this module's own generate_report_
# for_session (still just a wrapper, unchanged) and each of the four
# real mutation call sites: services/curation_report_service.py's
# get_or_create_report/regenerate_report, services/curation_chat_
# service.py's add_curation_chat_exchanges_to_report, and curation_
# chat.py's _accept_report_update (the easiest of the four to miss,
# since it's domain logic living outside either service module).

# Public (no leading underscore) since every one of the four real
# mutation call sites -- spread across services/curation_report_
# service.py, services/curation_chat_service.py, and curation_chat.py --
# references these by name rather than hardcoding the string literal
# independently in four places.
GENERATION_REASON_INITIAL = "initial"
GENERATION_REASON_REGENERATE = "regenerate"
GENERATION_REASON_CHAT_ADD_TO_REPORT = "chat_add_to_report"
GENERATION_REASON_CHAT_AUTO_UPDATE = "chat_auto_update"


def _utcnow_iso() -> str:
    """Factored out purely so a test can patch this one function
    (`patch.object(report, "_utcnow_iso", ...)`) for a deterministic
    created_at, without needing to patch the datetime module itself."""
    return datetime.now(timezone.utc).isoformat()


def append_report_version(session: PaperPoolSession, report: dict, generation_reason: str) -> dict:
    """report-quality Phase R3: the ONE shared mutation point every
    report-generating call site must go through. Appends `report`
    (already a complete, freshly-generated/regenerated report dict --
    report_template already stamped by generate_report/
    _regenerate_report_sections_with_sources) as a new version on
    `session.report_versions`, makes it the active version, and mirrors
    `session.report` to it -- the pre-R3 compatibility field every
    existing reader still uses, so nothing downstream of this call needs
    to know versioning exists at all.

    version_number is 1-indexed and simply len(session.report_versions)
    + 1 at call time -- no gaps, no reuse, since versions are only ever
    appended, never removed. version_id is a fresh uuid4 hex, same
    convention curation_chat.py's own exchange_id already uses.

    Never mutates an existing version -- old versions are immutable
    historical snapshots once appended (see PaperPoolSession.report_
    versions' own docstring): a later regeneration's citation-
    preservation/source-revocation rules only ever shape the NEW version
    being built here, never reach back and rewrite one already in the
    list.

    generation_reason is a plain descriptive string (this module's usual
    preference for a small, evolving vocabulary over a Literal/enum --
    see ReportTemplate's own docstring for the same reasoning applied to
    templates) -- the four values every current call site uses are
    "initial" (first-ever generation for a session), "regenerate"
    (whole-pool /report/regenerate), "chat_add_to_report" (curation-
    chat-add-to-report's selective regeneration), and "chat_auto_update"
    (chat's own accept-the-report-update-offer flow). Returns the
    newly-appended version dict itself, so a caller that also needs its
    version_id doesn't have to re-derive it.
    """
    version_id = uuid.uuid4().hex
    version = {
        "version_id": version_id,
        "version_number": len(session.report_versions) + 1,
        "created_at": _utcnow_iso(),
        "report_template": report.get("report_template", DEFAULT_REPORT_TEMPLATE),
        "generation_reason": generation_reason,
        "report": report,
    }
    session.report_versions.append(version)
    session.active_report_version_id = version_id
    session.report = report
    return version


def get_active_report_version(session: PaperPoolSession) -> dict | None:
    """Looks up session.active_report_version_id in session.report_
    versions. Returns None if the session has no report/versions yet at
    all (active_report_version_id is still None), or -- defensively --
    if it somehow doesn't resolve to anything in report_versions (should
    never happen in practice: append_report_version/
    activate_report_version below are the only two places that ever set
    active_report_version_id, and both always set it to a version_id
    simultaneously present in report_versions). Returning None rather
    than raising KeyError keeps this a safe read for callers like the
    API serializer, which already has its own graceful "no version
    metadata available" fallback for exactly this case."""
    if session.active_report_version_id is None:
        return None
    for version in session.report_versions:
        if version["version_id"] == session.active_report_version_id:
            return version
    return None


def activate_report_version(session: PaperPoolSession, version_id: str) -> dict | None:
    """report-quality Phase R3: switches session's active report version
    to `version_id`, mirroring session.report to that version's own
    report dict -- the same "session.report always mirrors whatever
    version is active" invariant append_report_version establishes.
    Never mutates the version list itself, only which entry is active
    and session.report's own pointer.

    Returns the now-active version dict, or None (mutates NOTHING) if
    version_id doesn't match anything in session.report_versions -- same
    "return None for absence rather than raise" convention curation_
    session.py's own load_curation_session already uses; the caller
    (services/curation_report_service.py) is what turns a None result
    into an HTTP 404, not this function itself.
    """
    for version in session.report_versions:
        if version["version_id"] == version_id:
            session.active_report_version_id = version_id
            session.report = version["report"]
            return version
    return None


# --- report-quality Phase R5A/R5C.1: document export of a report version ---
#
# Deterministic, backend-rendered, no LLM call -- walks a version's own
# already-finalized sections/references exactly as stored, never
# recomputes citations/references/markers. Deliberately excludes
# report["refinement"] (evaluator/refinement detail is internal QA
# information about HOW the report was produced, not part of the
# report's own content -- see R4.2's own "don't turn it into a
# dashboard" precedent, applied here to the exported artifact instead
# of the in-app panel) and chat_history/chat_references entirely (a
# report is a well-defined, versioned artifact; chat is a separate,
# unversioned scratchpad -- exporting "the report" must not silently
# also export "the chat").

_EXPORT_FILENAME_SANITIZE_RE = re.compile(r"[^a-zA-Z0-9]+")


@dataclass
class ExportSection:
    title: str
    # report-quality Phase R5C.1: split on the literal "\n\n" paragraph
    # separator, not a regex -- `"\n\n".join(content.split("\n\n"))` is
    # an EXACT inverse of the split for any string, so render_report_
    # markdown's own reconstruction below is guaranteed byte-identical
    # to the pre-R5C.1 implementation that just dumped `content` as-is,
    # regardless of how many blank lines a given report's prose happens
    # to contain. DOCX/PDF renderers get one paragraph per list entry
    # instead, which is the whole point of splitting at the model layer
    # rather than leaving it to each renderer to reinvent.
    paragraphs: list[str]


@dataclass
class ExportReference:
    number: int
    formatted: str
    link_url: str | None = None


@dataclass
class ReportExportDocument:
    """report-quality Phase R5C.1: the one logical document model every
    export renderer (Markdown today; DOCX as of this phase; PDF in a
    later phase) walks -- built once by build_report_export_document
    below, which is the ONLY place the legacy sections/references
    fallback, the title fallback, and the "created_at is optional" rule
    are decided. A renderer never re-derives any of that; it only knows
    how to lay out a ReportExportDocument in its own output format."""

    title: str
    # Ordered (label, value) pairs, not a dict -- display order matters
    # (Topic, Template, Version, then optional Generated) and a list of
    # tuples keeps that order explicit rather than relying on dict
    # insertion order being preserved correctly by every future caller.
    meta: list[tuple[str, str]]
    sections: list[ExportSection]
    references: list[ExportReference]


def build_report_export_document(session: PaperPoolSession, version: dict) -> ReportExportDocument:
    """report-quality Phase R5C.1: the single place that turns a stored
    report version into the document shape every export renderer
    consumes. Reuses the exact same "derive at read time" fallback
    convention api_app/serializers.py's _report_to_out already
    established for an old report with no `sections`/`references` key
    of its own yet -- calls derive_sections_from_legacy_report/
    derive_legacy_references directly rather than reimplementing that
    fallback a third time (now a SECOND time total, since this is the
    only remaining call site). Every section's `content` is used
    exactly as stored -- it already carries its final, resolved [N]
    markers (see _build_references_and_renumber, which runs once at
    generation/revision time, never again after) -- and every
    reference's `formatted` string is already a complete citation
    (format_apa_citation/format_web_citation, also computed once at
    generation time) -- this function only WALKS that already-finalized
    data into the document model; it never recomputes, re-numbers, or
    re-resolves any of it.
    """
    report = version["report"]
    if "references" not in report:
        report = derive_legacy_references(report)
    if "sections" not in report:
        report = {**report, "sections": derive_sections_from_legacy_report(report)}

    title = session.display_title or session.topic
    meta: list[tuple[str, str]] = [
        ("Topic", session.topic),
        ("Template", report.get("report_template", DEFAULT_REPORT_TEMPLATE).capitalize()),
        ("Version", f"{version['version_number']} ({version['generation_reason']})"),
    ]
    if version.get("created_at"):
        meta.append(("Generated", version["created_at"]))

    sections = [
        ExportSection(title=section["title"], paragraphs=section["content"].split("\n\n"))
        for section in report["sections"]
    ]
    references = [
        ExportReference(number=ref["number"], formatted=ref["formatted"], link_url=ref.get("link_url"))
        for ref in report.get("references", [])
    ]
    return ReportExportDocument(title=title, meta=meta, sections=sections, references=references)


def render_report_markdown(session: PaperPoolSession, version: dict) -> str:
    """report-quality Phase R5A, refactored in R5C.1 to consume the
    shared ReportExportDocument (built by build_report_export_document
    above) instead of walking session.report directly -- output is
    byte-identical to the pre-R5C.1 implementation (proven by the
    original R5A tests passing unchanged), since `"\\n\\n".join(section
    .paragraphs)` is an exact inverse of the `content.split("\\n\\n")`
    the document model used to build those paragraphs in the first
    place.

    References section is omitted entirely (not rendered with a
    "no references" placeholder) when the report has none -- the same
    choice the frontend's own ReferencesList already makes for the
    in-app view, kept consistent here rather than diverging for the
    exported document specifically.
    """
    doc = build_report_export_document(session, version)

    lines = [f"# {doc.title}", ""]
    for label, value in doc.meta:
        lines.append(f"**{label}:** {value}  ")
    lines.append("")

    for section in doc.sections:
        lines.append(f"## {section.title}")
        lines.append("")
        lines.append("\n\n".join(section.paragraphs))
        lines.append("")

    if doc.references:
        lines.append("## References")
        lines.append("")
        for ref in doc.references:
            citation = f"[{ref.formatted}]({ref.link_url})" if ref.link_url else ref.formatted
            lines.append(f"{ref.number}. {citation}")
        lines.append("")

    return "\n".join(lines)


def _add_docx_hyperlink(paragraph, url: str, text: str) -> None:
    """python-docx has no first-class add_hyperlink() -- this is the
    documented low-level workaround (build the w:hyperlink XML element
    by hand and register the URL as an external relationship on the
    paragraph's own part) every python-docx hyperlink recipe uses.
    Styled as a conventional blue/underlined link so it reads as a real
    hyperlink rather than plain text in the rendered document."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)  # noqa: SLF001 -- documented workaround, no public API exists


def render_report_docx(session: PaperPoolSession, version: dict) -> bytes:
    """report-quality Phase R5C.1: renders the same ReportExportDocument
    render_report_markdown consumes as a DOCX file, returned as raw
    bytes (the router wraps them in an HTTP Response; this function
    stays format-only, same "no HTTP knowledge" scope every other
    render_* function here already keeps).

    Layout: title (Heading 0) -> metadata lines (plain italic
    paragraphs, not a table -- a handful of label/value lines don't
    need one) -> one Heading 1 + one paragraph per non-empty content
    paragraph for each section, in order -> a page break, then a
    References heading and numbered entries (only when references
    exist -- same "omit entirely, don't render a placeholder" choice
    render_report_markdown already makes) with a real hyperlink
    wherever `link_url` is present, plain text otherwise. Deliberately
    no cover page, no custom fonts, no branding -- a clean, minimal
    document template, not a styled artifact.
    """
    doc = build_report_export_document(session, version)
    document = Document()

    document.add_heading(doc.title, level=0)
    for label, value in doc.meta:
        meta_paragraph = document.add_paragraph()
        meta_run = meta_paragraph.add_run(f"{label}: {value}")
        meta_run.italic = True

    for section in doc.sections:
        document.add_heading(section.title, level=1)
        for paragraph_text in section.paragraphs:
            if paragraph_text:
                document.add_paragraph(paragraph_text)

    if doc.references:
        document.add_page_break()
        document.add_heading("References", level=1)
        for ref in doc.references:
            ref_paragraph = document.add_paragraph()
            ref_paragraph.add_run(f"{ref.number}. ")
            if ref.link_url:
                _add_docx_hyperlink(ref_paragraph, ref.link_url, ref.formatted)
            else:
                ref_paragraph.add_run(ref.formatted)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_report_pdf(session: PaperPoolSession, version: dict) -> bytes:
    """report-quality Phase R5C.2: renders the same ReportExportDocument
    render_report_markdown/render_report_docx consume as a PDF, via
    ReportLab's Platypus flowables (SimpleDocTemplate + Paragraph/
    Spacer/PageBreak + the default style sheet) rather than raw canvas
    positioning -- gets a clean, document-like layout without hand-
    rolled coordinate math. Same "no HTTP knowledge, returns raw bytes"
    scope as render_report_docx.

    **Escaping is required, not optional.** ReportLab's Paragraph text
    is parsed as a small XML-like markup language (`<b>`, `<a href>`,
    etc.) -- every piece of text that ultimately comes from the report
    itself (title, metadata values, section titles, body paragraphs,
    reference "formatted" strings) is LLM- or user-authored and could
    structurally contain `<`, `>`, or `&`. Every one of those strings is
    run through xml_escape before being embedded in a Paragraph's
    source string; a hyperlink's URL is embedded via xml_quoteattr
    (not xml_escape -- an href value is an XML ATTRIBUTE, which needs
    quote-safe quoting, not text-node escaping) so a URL containing a
    literal `&` or a stray `"` can't break out of the `href="..."`
    attribute or otherwise malform the markup. Skipping this step is
    the one way this function could silently corrupt or crash on
    perfectly ordinary report content.

    Layout mirrors render_report_docx: title -> metadata (italic lines)
    -> one Heading1 + one Paragraph per non-empty content paragraph per
    section, in order -> a page break, then a References heading and
    numbered entries (only when references exist) with a real
    `<a href>` hyperlink wherever `link_url` is present. No cover page,
    no custom font embedding -- ReportLab's default Helvetica-based
    style sheet throughout.
    """
    doc = build_report_export_document(session, version)
    styles = getSampleStyleSheet()
    meta_style = ParagraphStyle("ExportMeta", parent=styles["Normal"], fontName="Helvetica-Oblique")
    reference_style = ParagraphStyle("ExportReference", parent=styles["Normal"], spaceAfter=6)

    flowables = [Paragraph(xml_escape(doc.title), styles["Title"]), Spacer(1, 12)]
    for label, value in doc.meta:
        flowables.append(Paragraph(f"{xml_escape(label)}: {xml_escape(value)}", meta_style))
    flowables.append(Spacer(1, 12))

    for section in doc.sections:
        flowables.append(Paragraph(xml_escape(section.title), styles["Heading1"]))
        for paragraph_text in section.paragraphs:
            if paragraph_text:
                flowables.append(Paragraph(xml_escape(paragraph_text), styles["BodyText"]))
        flowables.append(Spacer(1, 6))

    if doc.references:
        flowables.append(PageBreak())
        flowables.append(Paragraph("References", styles["Heading1"]))
        for ref in doc.references:
            escaped_citation = xml_escape(ref.formatted)
            if ref.link_url:
                escaped_citation = f'<a href={xml_quoteattr(ref.link_url)} color="blue">{escaped_citation}</a>'
            flowables.append(Paragraph(f"{ref.number}. {escaped_citation}", reference_style))

    buffer = io.BytesIO()
    SimpleDocTemplate(buffer, pagesize=letter).build(flowables)
    return buffer.getvalue()


def report_export_filename(session: PaperPoolSession, version: dict, extension: str) -> str:
    """report-quality Phase R5A: a safe, deterministic Content-
    Disposition filename -- sanitized display title (or topic, same
    fallback render_report_markdown's own title line already uses) plus
    version number, so re-exporting the same version always produces
    the identical filename, and a title containing spaces/punctuation/
    unicode can never produce an unsafe or malformed header value."""
    base = session.display_title or session.topic or "report"
    slug = _EXPORT_FILENAME_SANITIZE_RE.sub("-", base).strip("-").lower() or "report"
    return f"{slug}-v{version['version_number']}.{extension}"
